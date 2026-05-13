"""word_helpers.py
Funciones utilitarias para aplicar formato y bordes usando python‑docx.
"""

from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------------------- #
# PÁRRAFOS                                                               #
# ---------------------------------------------------------------------- #
def apply_paragraph_style(
    paragraph,
    *,
    font_name: str = "Arial",
    font_size: int = 7,
    bold: bool = False,
    alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
    line_spacing: float = 1.5,
    space_before: int = 0,
    space_after: int = 0,
    indent: float = 0.0,
):
    """Aplica propiedades a un párrafo y su primer run."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold

    paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.left_indent = Inches(indent)

    # Compatibilidad este/oeste
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    return paragraph


def add_paragraph_border(paragraph, position: str = "top"):
    """Agrega borde en posición {top|bottom} a un párrafo."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr")) or OxmlElement("w:pBdr")
    pPr.append(pBdr)
    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "4")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), "000000")
    pBdr.append(border)


def add_bullet_paragraph(document, text: str) -> None:
    """
    Inserta un párrafo con viñeta (‘List Bullet’) usando
    la misma tipografía y tamaño que el resto del documento (Arial 7 pt).
    """
    p = document.add_paragraph(text)
    # 1) convertirlo en viñeta
    if "List Bullet" in document.styles:          # plantilla latina de Word
        p.style = "List Bullet"
    else:                                         # fallback: estilo por defecto & símbolo
        p.style = document.styles["Normal"]
        p.runs[0].text = "\u2022 " + p.runs[0].text  # • resto del texto

    # 2) formatearlo igual que el cuerpo
    apply_paragraph_style(
        p,
        font_name="Arial",
        font_size=7,
        alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        line_spacing=1.5,
    )
    # 3) sangrías (½ cm primera línea negativa; 1 cm resto)
    pf = p.paragraph_format
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-0.5)
    pf.space_after = Pt(12)


# ---------------------------------------------------------------------- #
# TABLAS                                                                  #
# (todas las funciones que ya había: set_row_height, set_vertical_alignment,
#  set_cell_border)                                                       #
# ---------------------------------------------------------------------- #
def set_row_height(row, height_twips: int):
    tr = row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(height_twips))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def set_vertical_alignment(cell, alignment: str = "center"):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), alignment)
    tcPr.append(vAlign)


def set_cell_border(cell,
                    top=None, left=None, bottom=None, right=None,
                    insideH=None, insideV=None):
    """
    Aplica bordes al objeto cell (se deja tal cual estaba).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)

    def _set(side, attrs):
        if attrs is None:
            return
        side_el = borders.find(qn(f'w:{side}'))
        if side_el is None:
            side_el = OxmlElement(f'w:{side}')
            borders.append(side_el)
        for k, v in attrs.items():
            side_el.set(qn(f'w:{k}'), str(v))

    _set("top",      top)
    _set("left",     left)
    _set("bottom",   bottom)
    _set("right",    right)
    _set("insideH",  insideH)
    _set("insideV",  insideV)


# ---------------------------------------------------------------------- #
# CAMPOS DINÁMICOS (add_page_number) – sin cambios                       #
# ---------------------------------------------------------------------- #
def add_page_number(paragraph):
    """Inserta "Página X de Y" en el párrafo dado."""
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.extend([fldBegin, instr, fldEnd])

    run.add_text(" de ")

    run2 = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "NUMPAGES"
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run2._r.extend([fldBegin, instr, fldEnd])
