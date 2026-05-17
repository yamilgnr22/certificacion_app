# document_generator.py
"""
Orquesta la creación del .docx completo
(Secciones + inserción de plantilla SmartArt al final)
"""

from __future__ import annotations
import os
import tempfile
from typing import Optional

from docx import Document
from docxcompose.composer import Composer

from generators.certificacion import generar_certificacion
from generators.er_table      import generar_tabla_er
from generators.esf_table     import generar_tabla_esf
from generators.esf_table_mensual import generar_tabla_esf_mensual
from generators.datos_table   import generar_tabla_datos
from generators.docs_table    import generar_tabla_docs_cliente   # ⭐️ nuevo
from validators               import validate_er, validate_esf

# -------------- Ruta por defecto a la plantilla SmartArt --------------
_PLANTILLA_PATH = os.path.join(
    os.path.dirname(__file__),
    "plantilla_smartArt.docx",
)

# -------------- Fusión con la plantilla SmartArt ----------------------
def _fusionar_con_plantilla(doc_final: Document, ruta_plantilla: str, salida: str):
    composer = Composer(doc_final)
    composer.append(Document(ruta_plantilla))
    composer.save(salida)

# -------------- API pública -------------------------------------------
def generar_documento_completo(
    df_esf,
    df_er,
    df_datos,
    df_cert,
    output_path: str,
    plantilla_path: Optional[str] = None,
    incluir_validacion: bool = True,
    tolerancia_validacion: float = 1.0,
    detener_si_error: bool = False,
    validacion_documentos: Optional[dict] = None,
    validacion_llm: Optional[dict] = None,
    statement_blocks: Optional[list] = None,
    esf_tipo: str = "corte",
):
    """
    Genera el documento final y lo guarda en `output_path`.
    Orden de secciones:
        1) Certificación
        2) Estado de Resultados
        3) Estado de Situación Financiera
        4) Datos
        5) Documentos del cliente (tabla vacía)
        6) Plantilla SmartArt (última página)
    """
    doc = Document()

    # -------- Validación de consistencia (opcional) --------
    if incluir_validacion:
        v_er  = validate_er(df_er, tolerance=tolerancia_validacion)
        # Validación acorde al tipo de ESF
        v_esf = validate_esf(df_esf, tolerance=tolerancia_validacion, mode=(esf_tipo or "corte"))

        titulo = doc.add_paragraph("Validación de consistencia")
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from word_helpers import apply_paragraph_style
        apply_paragraph_style(
            titulo, font_name="Arial", font_size=12, bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, line_spacing=1.15
        )

        def _add_line(text: str, ok: bool | None = None):
            p = doc.add_paragraph(text)
            apply_paragraph_style(p, font_name="Arial", font_size=9)
            if ok is not None:
                # Prefijo simple OK/ERROR
                p.runs[0].text = ("✅ " if ok else "❌ ") + p.runs[0].text

        _add_line("Estado de Resultados:")
        if v_er["checks"]:
            for c in v_er["checks"]:
                rule = c.get("rule", "Regla")
                col  = c.get("column")
                ok   = c.get("ok")
                _add_line(f"{rule} [col {col}]", ok)
        if v_er["errors"]:
            for e in v_er["errors"]:
                _add_line(e, False)

        _add_line("")
        _add_line("Estado de Situación Financiera:")
        if v_esf["checks"]:
            for c in v_esf["checks"]:
                rule = c.get("rule", "Regla")
                ok   = c.get("ok")
                _add_line(f"{rule}", ok)
        if v_esf["errors"]:
            for e in v_esf["errors"]:
                _add_line(e, False)

        doc.add_page_break()

        if detener_si_error and (not v_er.get("ok", True) or not v_esf.get("ok", True)):
            # Resumen en excepción (también queda plasmado en la página de validación)
            errs = [*v_er.get("errors", []), *v_esf.get("errors", [])]
            raise ValueError("Validación contable fallida: " + "; ".join(errs))

    # -------- Validación de cédula por visión (mostrar en DOCX sólo si incluir_validacion=True) --------
    if incluir_validacion and (validacion_documentos is not None):
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from word_helpers import apply_paragraph_style
        titulo = doc.add_paragraph("Validación de Cédula (Visión)")
        apply_paragraph_style(
            titulo, font_name="Arial", font_size=12, bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, line_spacing=1.15
        )
        ok_global = validacion_documentos.get("ok", True)
        p = doc.add_paragraph("Resultado global: ")
        apply_paragraph_style(p, font_name="Arial", font_size=9)
        p.add_run("OK" if ok_global else "CON INCIDENTES").bold = True
        # Detalle
        for chk in validacion_documentos.get("checks", []):
            ok = chk.get("ok", False)
            campo = chk.get("field", "campo")
            docname = chk.get("doc", "doc")
            exp = chk.get("expected")
            got = chk.get("got")
            line = f"[{docname}] {campo}: esperado=‘{exp}’, obtenido=‘{got}’"
            q = doc.add_paragraph(line)
            apply_paragraph_style(q, font_name="Arial", font_size=9)
            # prefijo visual
            q.runs[0].text = ("✅ " if ok else "❌ ") + q.runs[0].text
        doc.add_page_break()

    # -------- Validación LLM (mostrar en DOCX sólo si incluir_validacion=True) --------
    if incluir_validacion and (validacion_llm is not None):
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from word_helpers import apply_paragraph_style
        titulo = doc.add_paragraph("Validación LLM")
        apply_paragraph_style(
            titulo, font_name="Arial", font_size=12, bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, line_spacing=1.15
        )
        p = doc.add_paragraph(validacion_llm.get("summary", ""))
        apply_paragraph_style(p, font_name="Arial", font_size=9)
        issues = validacion_llm.get("issues", [])
        for it in issues:
            sev = it.get("severity", "")
            desc = it.get("description", "")
            ev  = it.get("evidence", "")
            sug = it.get("suggestion", "")
            line = f"[{sev}] {desc}\nEvidencia: {ev}\nSugerencia: {sug}"
            q = doc.add_paragraph(line)
            apply_paragraph_style(q, font_name="Arial", font_size=9)
        doc.add_page_break()

    generar_certificacion(doc, df_cert)

    blocks = statement_blocks or []
    if blocks:
        for block in blocks:
            block_cert = block.get("df_certificacion")
            if block_cert is None:
                block_cert = df_cert
            generar_tabla_er(doc, block.get("df_er"), block_cert)
            generar_tabla_esf_mensual(doc, block.get("df_esf_mensual"), block_cert)
    else:
        generar_tabla_er(doc,  df_er,  df_cert)
        # ESF según tipo
        if (esf_tipo or "corte").lower() == "mensual":
            generar_tabla_esf_mensual(doc, df_esf, df_cert)
        else:
            generar_tabla_esf(doc, df_esf, df_cert)
    generar_tabla_datos(doc, df_datos)

    # -------- NUEVA SECCIÓN --------
    generar_tabla_docs_cliente(doc)          # ← aquí insertamos la tabla vacía

    # -------- Plantilla SmartArt --------
    tpl = plantilla_path or _PLANTILLA_PATH
    if not os.path.isfile(tpl):
        raise FileNotFoundError(f"Plantilla SmartArt no encontrada: {tpl}")

    # Guardar en temporal para fusionar
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_name = tmp.name
    doc.save(tmp_name)

    _fusionar_con_plantilla(Document(tmp_name), tpl, output_path)
    os.remove(tmp_name)

