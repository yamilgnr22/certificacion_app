from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from document_generator import generar_documento_completo
from financial_model import build_financial_model
from tests.test_financial_model import _docx_text, month_column, sample_payload


def _fmt(value: float) -> str:
    return f"{float(value):,.0f}"


class DocumentE2ETest(unittest.TestCase):
    """F5-T2: el DOCX generado refleja las cifras del modelo, no otras."""

    @classmethod
    def setUpClass(cls):
        cls.result = build_financial_model(sample_payload())
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            cls.out_path = Path(tmp.name)
        generar_documento_completo(
            cls.result.df_esf_mensual,
            cls.result.df_er,
            cls.result.df_datos,
            cls.result.df_certificacion,
            str(cls.out_path),
            incluir_validacion=False,
            detener_si_error=False,
            esf_tipo="mensual",
        )
        cls.text = _docx_text(cls.out_path)

    @classmethod
    def tearDownClass(cls):
        cls.out_path.unlink(missing_ok=True)

    def test_client_identity_present(self):
        self.assertIn("Cliente Prueba", self.text)
        self.assertIn("001-010101-0000A", self.text)

    def test_certification_amounts_match_model_summary(self):
        summary = self.result.summary
        # Ingresos brutos del periodo y promedio mensual (texto legal).
        self.assertIn(_fmt(summary["income_total"]), self.text)
        self.assertIn(_fmt(summary["income_average"]), self.text)
        # Utilidad del periodo y promedio.
        self.assertIn(_fmt(summary["net_income_total"]), self.text)
        self.assertIn(_fmt(summary["net_income_average"]), self.text)

    def test_er_table_totals_match_dataframe(self):
        df = self.result.df_er
        total_row = df[df["Descripcion"] == "Total gastos operativos"].iloc[0]
        accumulated = float(total_row["Acumulado del periodo"])
        self.assertIn(_fmt(accumulated), self.text)
        # El total del primer mes tambien debe estar en la tabla.
        first_month_col = df.columns[2]
        self.assertIn(_fmt(float(total_row[first_month_col])), self.text)

    def test_esf_table_totals_match_dataframe(self):
        df = self.result.df_esf_mensual
        last_month = self.result.summary["months"][-1]
        col = month_column(df, last_month)
        for label in ["Total Activos", "Total Pasivos", "Total Patrimonio", "Total Pasivo + Patrimonio"]:
            row = df[df["Descripcion"] == label].iloc[0]
            self.assertIn(
                _fmt(float(row[col])),
                self.text,
                f"el DOCX no contiene el valor de '{label}' del ultimo mes",
            )

    def test_document_has_expected_sections(self):
        self.assertIn("CERTIFICACIÓN DEL CONTADOR PÚBLICO INDEPENDIENTE", self.text)
        self.assertIn("Estado de Resultados", self.text)
        self.assertIn("Estado de Situación Financiera", self.text)

    def test_docx_opens_with_tables(self):
        doc = Document(str(self.out_path))
        self.assertGreaterEqual(len(doc.tables), 2)


class CertificadoEnterosTest(unittest.TestCase):
    """El certificado se expresa en cordobas enteros: la cifra y las letras
    coinciden y nunca aparece 'punto' en la conversion a palabras (bug del
    caso Jose David: promedios con decimales daban '...punto tres tres')."""

    def _texto_certificado(self, ingresos_prom, utilidad_prom) -> str:
        import pandas as pd
        from docx import Document
        from generators.certificacion import generar_certificacion

        dfc = pd.DataFrame([
            ["Nombre completo", "Cliente Decimal"],
            ["cedula", "001-010101-0000A"],
            ["inicio", "2026-01-01"],
            ["fin", "2026-06-30"],
            ["sexo", "masculino"],
            ["banco", "Banco X"],
            ["Fecha Certificacion", "2026-07-05"],
            ["ingresos brutos", 23657828.00],
            ["promedio ingresos", ingresos_prom],
            ["utilidad del periodo", 2833355.35],
            ["promedio utilidad", utilidad_prom],
        ])
        doc = Document()
        generar_certificacion(doc, dfc)
        return "\n".join(p.text for p in doc.paragraphs)

    def test_promedios_decimales_no_generan_punto(self):
        texto = self._texto_certificado(3942971.33, 472226.89)
        # Ninguna cantidad en letras debe contener 'punto' (decimales).
        for linea in texto.split("\n"):
            if "córdobas" in linea:
                self.assertNotIn("punto", linea, f"decimales en letras: {linea!r}")

    def test_cifra_y_letra_coinciden_redondeadas(self):
        # 472,226.89 redondea a 472,227 -> la cifra y la letra deben coincidir.
        texto = self._texto_certificado(3942971.33, 472226.89)
        self.assertIn("472,227", texto)
        self.assertIn("cuatrocientos setenta y dos mil doscientos veintisiete", texto)
        self.assertNotIn("472,226", texto)


if __name__ == "__main__":
    unittest.main()
