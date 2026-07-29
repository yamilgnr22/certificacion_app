"""Tests de persistencia del Motor V2 (periodos engine='v2').

Ciclo completo contra la API con sqlite en memoria:
  crear borrador -> listar -> actualizar -> finalizar (DOCX + saldos cacheados)
  -> rollforward propuesto -> guards del editor clasico (no puede tocar v2).
"""

from __future__ import annotations

import json
import os
import tempfile
import unittest
from pathlib import Path

from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

import web_server
from db.models import Base
from db.seed import seed_giros


MESES = ["2025-12", "2026-01", "2026-02", "2026-03", "2026-04", "2026-05"]
ING = [332915, 543402, 301554, 325063, 298999, 425647]
COGS = [166524, 277026, 152888, 164774, 150127, 213888]


def gloria_inputs() -> dict:
    er = [
        {
            "mes": m, "ingresos": ING[i], "costo_ventas": COGS[i],
            "sueldos_salarios": 15016, "servicios_publicos": 5311, "alcaldia_dgi": 732,
            "combustible": 5311, "publicidad": 5494, "renta": 9522,
            "gasto_depreciacion": 9003, "otros_gastos": 1831,
        }
        for i, m in enumerate(MESES)
    ]
    return {
        "periodo": {"tipo": "A", "mes_inicial": "2025-12", "mes_final": "2026-05", "tasa_cambio": 36.6243},
        "datos": {
            "nombre_completo": "Gloria Elena Guillen Robinson", "cedula": "601-140998-0002L",
            "empleados": 1, "banco": "FICOHSA", "fecha_certificacion": "2026-06-05",
            "regimen": "Cuota Fija",
        },
        "er_mensual": er,
        "saldos_iniciales": {"mobiliario_equipos": 366243, "vehiculos": 695862, "tarjetas_credito": 62261},
        "saldos_finales": {
            "efectivo": 841220, "mobiliario_equipos": 366243, "vehiculos": 695862,
            "depreciacion_acumulada": -54018, "tarjetas_credito": 60430,
        },
        "deudas": [{
            "numero": "TARJETA", "entidad": "Tarjeta", "tipo_credito": "TARJETA DE CREDITO",
            "estrategia": "revolving", "moneda": "NIO", "valor_inicial": 62261,
            "saldo_reportado": 60430, "cuota": 0, "fecha_otorgamiento": "2020-01-01",
            "saldo_apertura": 62261, "incluir_en_er": True,
            "saldos_mensuales": {
                "2025-12": 40287, "2026-01": 38456, "2026-02": 43949,
                "2026-03": 45780, "2026-04": 47612, "2026-05": 60430,
            },
        }],
    }


class MotorV2PersistenciaTest(unittest.TestCase):
    def setUp(self):
        self.engine = create_engine("sqlite:///:memory:", future=True)
        Base.metadata.create_all(self.engine)
        self.factory = sessionmaker(bind=self.engine, expire_on_commit=False, future=True)
        with self.factory() as session:
            seed_giros(session)
            session.commit()
        self.old_engine = web_server.app.config.get("DB_ENGINE")
        self.old_require = web_server.app.config.get("DB_REQUIRE_ALEMBIC")
        web_server.app.config["DB_ENGINE"] = self.engine
        web_server.app.config["DB_REQUIRE_ALEMBIC"] = False
        self.client = web_server.app.test_client()
        self.tmpdir = tempfile.TemporaryDirectory()
        self.old_docs_dir = os.environ.get("CERTAPP_DOCUMENTOS_DIR")
        os.environ["CERTAPP_DOCUMENTOS_DIR"] = self.tmpdir.name
        self.old_soportes_dir = os.environ.get("CERTAPP_SOPORTES_DIR")
        os.environ["CERTAPP_SOPORTES_DIR"] = self.tmpdir.name

    def tearDown(self):
        if self.old_engine is None:
            web_server.app.config.pop("DB_ENGINE", None)
        else:
            web_server.app.config["DB_ENGINE"] = self.old_engine
        if self.old_require is None:
            web_server.app.config.pop("DB_REQUIRE_ALEMBIC", None)
        else:
            web_server.app.config["DB_REQUIRE_ALEMBIC"] = self.old_require
        if self.old_docs_dir is None:
            os.environ.pop("CERTAPP_DOCUMENTOS_DIR", None)
        else:
            os.environ["CERTAPP_DOCUMENTOS_DIR"] = self.old_docs_dir
        if self.old_soportes_dir is None:
            os.environ.pop("CERTAPP_SOPORTES_DIR", None)
        else:
            os.environ["CERTAPP_SOPORTES_DIR"] = self.old_soportes_dir
        self.tmpdir.cleanup()

    # ------------------------------------------------------------ helpers
    def crear_cliente(self) -> str:
        resp = self.client.post("/api/clientes", json={
            "nombre_completo": "Gloria Elena Guillen Robinson",
            "cedula": "601-140998-0002L",
            "nombre_negocio": "Envios GG",
            "direccion_negocio": "Bolonia",
            "giro_negocio_id": "ferreteria",
        })
        self.assertEqual(resp.status_code, 201, resp.get_json())
        return resp.get_json()["cliente"]["id"]

    def crear_borrador(self, cliente_id: str, inputs: dict | None = None) -> dict:
        resp = self.client.post("/api/motor/v2/periodos", json={
            "cliente_id": cliente_id,
            "inputs": inputs or gloria_inputs(),
        })
        self.assertEqual(resp.status_code, 201, resp.get_json())
        return resp.get_json()["periodo"]

    # -------------------------------------------------------------- tests
    def test_crear_borrador_y_listar(self):
        cid = self.crear_cliente()
        periodo = self.crear_borrador(cid)
        self.assertEqual(periodo["estado"], "borrador")
        self.assertEqual(periodo["engine"], "v2")
        self.assertEqual(periodo["mes_inicial"], "2025-12")
        self.assertEqual(periodo["periodo_meses"], 6)

        resp = self.client.get(f"/api/motor/v2/periodos?cliente_id={cid}")
        periodos = resp.get_json()["periodos"]
        self.assertEqual(len(periodos), 1)
        self.assertEqual(periodos[0]["id"], periodo["id"])

    def test_inputs_invalidos_rechazados(self):
        cid = self.crear_cliente()
        malos = gloria_inputs()
        malos["periodo"]["mes_final"] = "2025-01"  # anterior al inicial
        resp = self.client.post("/api/motor/v2/periodos", json={"cliente_id": cid, "inputs": malos})
        self.assertEqual(resp.status_code, 400, resp.get_json())

    def test_actualizar_persiste_inputs(self):
        cid = self.crear_cliente()
        periodo = self.crear_borrador(cid)
        nuevos = gloria_inputs()
        nuevos["er_mensual"][0]["ingresos"] = 999999
        resp = self.client.put(f"/api/motor/v2/periodos/{periodo['id']}", json={"inputs": nuevos})
        self.assertEqual(resp.status_code, 200, resp.get_json())

        resp = self.client.get(f"/api/motor/v2/periodos/{periodo['id']}")
        inputs = resp.get_json()["inputs"]
        self.assertEqual(inputs["er_mensual"][0]["ingresos"], 999999)

    def test_editor_clasico_no_puede_tocar_v2(self):
        cid = self.crear_cliente()
        periodo = self.crear_borrador(cid)
        pid = periodo["id"]

        r_update = self.client.put(f"/api/periodos/{pid}", json={"tasa_cambio": 37.0})
        self.assertEqual(r_update.status_code, 409, r_update.get_json())
        r_preview = self.client.post(f"/api/periodos/{pid}/preview")
        self.assertEqual(r_preview.status_code, 409, r_preview.get_json())
        r_fin = self.client.post(f"/api/periodos/{pid}/finalizar")
        self.assertEqual(r_fin.status_code, 409, r_fin.get_json())

        r_editables = self.client.get("/api/periodos/editables")
        ids = [p["id"] for p in r_editables.get_json().get("periodos", [])]
        self.assertNotIn(pid, ids)

    def test_finalizar_feliz_guarda_docx_y_saldos(self):
        cid = self.crear_cliente()
        periodo = self.crear_borrador(cid)
        resp = self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar")
        data = resp.get_json()
        self.assertEqual(resp.status_code, 200, data)
        self.assertTrue(data["ok"])
        self.assertEqual(data["periodo"]["estado"], "finalizado")
        self.assertTrue(Path(data["documento_path"]).exists())
        self.assertGreater(Path(data["documento_path"]).stat().st_size, 10_000)

        # Documento descargable
        r_doc = self.client.get(f"/api/motor/v2/periodos/{periodo['id']}/documento")
        self.assertEqual(r_doc.status_code, 200)

        # Saldos finales cacheados con las cifras de Gloria
        with self.factory() as session:
            from db.models import PeriodoCertificacion
            row = session.get(PeriodoCertificacion, periodo["id"])
            saldos = json.loads(row.saldos_finales_json)
        self.assertAlmostEqual(saldos["efectivo"], 841220.0, places=1)
        self.assertAlmostEqual(saldos["capital"], 999844.0, places=1)
        self.assertAlmostEqual(saldos["resultados_ejercicio"], 789033.0, places=1)

        # No editable tras finalizar
        r_upd = self.client.put(
            f"/api/motor/v2/periodos/{periodo['id']}", json={"inputs": gloria_inputs()}
        )
        self.assertEqual(r_upd.status_code, 409, r_upd.get_json())

    def test_finalizar_con_esf_vista_mensual(self):
        # esf_vista='mensual' persiste en el borrador y el DOCX final sale
        # con el ESF por meses (tabla ancha) en vez de la foto al corte.
        cid = self.crear_cliente()
        inputs = gloria_inputs()
        inputs["esf_vista"] = "mensual"
        periodo = self.crear_borrador(cid, inputs)

        r_get = self.client.get(f"/api/motor/v2/periodos/{periodo['id']}")
        self.assertEqual(r_get.get_json()["inputs"].get("esf_vista"), "mensual")

        resp = self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar")
        data = resp.get_json()
        self.assertEqual(resp.status_code, 200, data)
        from docx import Document
        doc = Document(data["documento_path"])
        # El ESF mensual tiene una columna por mes (desc + 6 meses = 7)
        anchas = [t for t in doc.tables if len(t.columns) == 7 and any(
            "Efectivo" in r.cells[0].text for r in t.rows
        )]
        self.assertEqual(len(anchas), 1, "no se encontro la tabla ESF mensual de 7 columnas")

    def test_finalizar_con_descuadre_da_422_y_sigue_borrador(self):
        cid = self.crear_cliente()
        inputs = gloria_inputs()
        inputs["saldos_finales"]["efectivo"] = 111111  # no cuadra con el modelo
        periodo = self.crear_borrador(cid, inputs)
        resp = self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar")
        self.assertEqual(resp.status_code, 422, resp.get_json())
        data = resp.get_json()
        self.assertFalse(data["ok"])
        self.assertTrue(data["validacion"]["errores"])

        r_get = self.client.get(f"/api/motor/v2/periodos/{periodo['id']}")
        self.assertEqual(r_get.get_json()["periodo"]["estado"], "borrador")

    def test_saldos_rollforward_desde_v2_finalizado(self):
        cid = self.crear_cliente()
        periodo = self.crear_borrador(cid)
        self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar")

        resp = self.client.get(f"/api/motor/v2/clientes/{cid}/saldos-rollforward")
        data = resp.get_json()
        self.assertTrue(data["has_anterior"])
        saldos = data["saldos"]
        self.assertAlmostEqual(saldos["efectivo"], 841220.0, places=1)
        # RA nuevo = RA corte (0) + RE corte (789,033)
        self.assertAlmostEqual(saldos["resultados_acumulados"], 789033.0, places=1)
        # No traslada capital ni resultados del ejercicio como cuentas
        self.assertNotIn("capital", saldos)
        self.assertNotIn("resultados_ejercicio", saldos)

    def test_rollforward_sin_periodos_finalizados(self):
        cid = self.crear_cliente()
        resp = self.client.get(f"/api/motor/v2/clientes/{cid}/saldos-rollforward")
        self.assertFalse(resp.get_json()["has_anterior"])

    # -------------------------------------- documentos soporte (imagenes)
    _PNG_1PX = __import__("base64").b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
    )

    def subir_documento(
        self, cid: str, tipo: str = "cedula_front", nombre: str = "cedula.png",
        contenido: bytes | None = None,
    ) -> dict:
        import io
        resp = self.client.post(
            f"/api/clientes/{cid}/documentos",
            data={"archivo": (io.BytesIO(contenido or self._PNG_1PX), nombre), "tipo": tipo},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 201, resp.get_json())
        return resp.get_json()["documento"]

    def test_documentos_subir_listar_servir_eliminar(self):
        cid = self.crear_cliente()
        doc = self.subir_documento(cid, "matricula", "matricula2026.png")
        self.assertEqual(doc["tipo"], "matricula")

        lista = self.client.get(f"/api/clientes/{cid}/documentos").get_json()["documentos"]
        self.assertEqual(len(lista), 1)
        self.assertEqual(lista[0]["original_filename"], "matricula2026.png")

        r_img = self.client.get(f"/api/documentos/{doc['id']}/archivo")
        self.assertEqual(r_img.status_code, 200)
        self.assertEqual(r_img.data, self._PNG_1PX)
        r_img.close()  # Windows: liberar el handle del archivo antes de borrar

        self.assertEqual(self.client.delete(f"/api/documentos/{doc['id']}").status_code, 200)
        lista2 = self.client.get(f"/api/clientes/{cid}/documentos").get_json()["documentos"]
        self.assertEqual(lista2, [])
        self.assertEqual(self.client.get(f"/api/documentos/{doc['id']}/archivo").status_code, 404)

    def test_documento_cambiar_tipo(self):
        """Arrastrar una imagen a otra casilla de la hoja: antes habia que
        borrarla de la biblioteca y volver a subirla."""
        cid = self.crear_cliente()
        doc = self.subir_documento(cid, "matricula_2", "matricula_2024.png")

        r = self.client.patch(f"/api/documentos/{doc['id']}", json={"tipo": "soporte_2"})
        self.assertEqual(r.status_code, 200, r.get_json())
        self.assertEqual(r.get_json()["documento"]["tipo"], "soporte_2")

        lista = self.client.get(f"/api/clientes/{cid}/documentos").get_json()["documentos"]
        self.assertEqual(len(lista), 1, "no se duplica: es la misma imagen")
        self.assertEqual(lista[0]["tipo"], "soporte_2")
        self.assertEqual(lista[0]["original_filename"], "matricula_2024.png")

    def test_documento_cambiar_tipo_invalido_es_400(self):
        cid = self.crear_cliente()
        doc = self.subir_documento(cid)
        r = self.client.patch(f"/api/documentos/{doc['id']}", json={"tipo": "matricula_9"})
        self.assertEqual(r.status_code, 400)
        sigue = self.client.get(f"/api/clientes/{cid}/documentos").get_json()["documentos"][0]
        self.assertEqual(sigue["tipo"], "cedula_front", "el tipo original no se toca")

    def test_documento_cambiar_tipo_inexistente_es_404(self):
        r = self.client.patch("/api/documentos/no-existe", json={"tipo": "cedula_back"})
        self.assertEqual(r.status_code, 404)

    def test_documento_tipo_invalido_rechazado(self):
        import io
        cid = self.crear_cliente()
        resp = self.client.post(
            f"/api/clientes/{cid}/documentos",
            data={"archivo": (io.BytesIO(self._PNG_1PX), "x.png"), "tipo": "pasaporte"},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 400, resp.get_json())

    def test_finalizar_incrusta_imagenes_seleccionadas(self):
        cid = self.crear_cliente()
        d1 = self.subir_documento(cid, "cedula_front", "ced_frente.png")
        d2 = self.subir_documento(cid, "matricula", "matricula.png")
        inputs = gloria_inputs()
        inputs["documentos_ids"] = [d1["id"], d2["id"]]
        periodo = self.crear_borrador(cid, inputs)

        # documentos_ids persistidos con el borrador
        guardado = self.client.get(f"/api/motor/v2/periodos/{periodo['id']}").get_json()["inputs"]
        self.assertEqual(guardado.get("documentos_ids"), [d1["id"], d2["id"]])

        data = self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar").get_json()
        self.assertTrue(data["ok"], data)
        from docx import Document
        doc = Document(data["documento_path"])
        tipos = {p.content_type for p in doc.part.package.iter_parts()}
        self.assertTrue(any(t.startswith("image/png") for t in tipos),
                        "el DOCX no contiene las imagenes incrustadas")

    def test_finalizar_sin_documentos_mantiene_tabla_vacia(self):
        cid = self.crear_cliente()
        periodo = self.crear_borrador(cid)  # sin documentos_ids
        data = self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar").get_json()
        self.assertTrue(data["ok"], data)
        from docx import Document
        doc = Document(data["documento_path"])
        tipos = {p.content_type for p in doc.part.package.iter_parts()}
        self.assertFalse(any(t.startswith("image/png") for t in tipos))

    def test_finalizar_notas_apagadas_docs_siempre_van(self):
        # Las notas son opcionales; los documentos del cliente SIEMPRE se
        # imprimen (aunque un borrador viejo traiga el flag legado incluir_docs).
        cid = self.crear_cliente()
        d1 = self.subir_documento(cid, "cedula_front", "ced_frente.png")
        inputs = gloria_inputs()
        inputs["documentos_ids"] = [d1["id"]]
        inputs["incluir_docs"] = False  # flag legado: se ignora
        inputs["incluir_notas"] = False
        periodo = self.crear_borrador(cid, inputs)

        data = self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar").get_json()
        self.assertTrue(data["ok"], data)
        from docx import Document
        doc = Document(data["documento_path"])
        tipos = {p.content_type for p in doc.part.package.iter_parts()}
        self.assertTrue(any(t.startswith("image/png") for t in tipos),
                        "los documentos del cliente siempre se incrustan")
        textos = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Documentos del cliente", textos)
        self.assertNotIn("Notas a los Estados Financieros", textos)

    def test_finalizar_hojas_default_incluidas(self):
        # Sin banderas (borradores viejos): docs + notas van; fotos del
        # negocio NO (es opt-in).
        cid = self.crear_cliente()
        periodo = self.crear_borrador(cid)
        data = self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar").get_json()
        self.assertTrue(data["ok"], data)
        from docx import Document
        doc = Document(data["documento_path"])
        textos = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Documentos del cliente", textos)
        self.assertIn("Notas a los Estados Financieros", textos)
        # NINGUNA variante de hoja de fotos (ni la nuestra ni la de la
        # plantilla SmartArt, que dice "Fotografias del negocio" sin tilde).
        self.assertNotIn("fotograf", textos.lower())

    def test_finalizar_fotos_negocio_sin_fotos_usa_plantilla(self):
        # Check ON pero sin fotos cargadas: sale la hoja de la plantilla
        # SmartArt ("Fotografias del negocio" con placeholders manuales),
        # una sola vez.
        cid = self.crear_cliente()
        inputs = gloria_inputs()
        inputs["incluir_fotos_negocio"] = True
        periodo = self.crear_borrador(cid, inputs)
        data = self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar").get_json()
        self.assertTrue(data["ok"], data)
        from docx import Document
        doc = Document(data["documento_path"])
        textos = "\n".join(p.text for p in doc.paragraphs)
        self.assertEqual(textos.lower().count("fotograf"), 1)

    def test_finalizar_con_cuentas_bancarias_en_nota1(self):
        # Nota 1 desglosada: cuenta bancaria + Efectivo en Caja residuo en el DOCX.
        cid = self.crear_cliente()
        inputs = gloria_inputs()
        inputs["cuentas_bancarias"] = [
            {"banco": "LAFISE", "tipo": "Cuenta de Ahorro", "moneda": "NIO",
             "numero": "106012140", "saldo": 500_000},
        ]
        periodo = self.crear_borrador(cid, inputs)
        data = self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar").get_json()
        self.assertTrue(data["ok"], data)
        from docx import Document
        doc = Document(data["documento_path"])
        celdas = {c.text for tb in doc.tables for r in tb.rows for c in r.cells}
        self.assertIn("LAFISE Cuenta de Ahorro NIO No. 106012140", celdas)
        self.assertIn("Efectivo en Caja", celdas)
        self.assertIn("341,220", celdas)  # 841,220 - 500,000 (residuo)

    def test_finalizar_caja_negativa_bloquea(self):
        # Cuentas suman mas que el efectivo del ESF -> NO finaliza (borrador intacto).
        cid = self.crear_cliente()
        inputs = gloria_inputs()
        inputs["cuentas_bancarias"] = [
            {"banco": "LAFISE", "tipo": "Cuenta de Ahorro", "moneda": "NIO",
             "numero": "1", "saldo": 900_000},  # > 841,220
        ]
        periodo = self.crear_borrador(cid, inputs)
        data = self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar").get_json()
        self.assertFalse(data["ok"])
        self.assertIn("negativo", data["validacion"]["errores"][0]["mensaje"])
        sigue = self.client.get(f"/api/motor/v2/periodos/{periodo['id']}").get_json()["periodo"]
        self.assertEqual(sigue["estado"], "borrador")  # no se finalizo

    def test_finalizar_con_fotos_negocio(self):
        # incluir_fotos_negocio=True agrega la hoja; las imagenes tipo
        # foto_negocio van ahi (y no al pareado de documentos).
        import io as _io
        from PIL import Image

        buf = _io.BytesIO()
        Image.new("RGB", (2, 2), (255, 0, 0)).save(buf, "PNG")  # distinta a _PNG_1PX
        cid = self.crear_cliente()
        d1 = self.subir_documento(cid, "cedula_front", "ced.png")
        d2 = self.subir_documento(cid, "foto_negocio", "local.png", contenido=buf.getvalue())
        inputs = gloria_inputs()
        inputs["documentos_ids"] = [d1["id"], d2["id"]]
        inputs["incluir_fotos_negocio"] = True
        periodo = self.crear_borrador(cid, inputs)

        data = self.client.post(f"/api/motor/v2/periodos/{periodo['id']}/finalizar").get_json()
        self.assertTrue(data["ok"], data)
        from docx import Document
        doc = Document(data["documento_path"])
        textos = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Documentos del cliente", textos)
        self.assertIn("Fotografías del Negocio", textos)
        tipos = [p.content_type for p in doc.part.package.iter_parts()]
        self.assertEqual(sum(1 for t in tipos if t.startswith("image/png")), 2)
        # UNA sola hoja de fotos: con fotos cargadas no se fusiona ademas la
        # plantilla SmartArt (que trae su propia "Fotografias del negocio").
        self.assertEqual(textos.lower().count("fotograf"), 1)


if __name__ == "__main__":
    unittest.main()
