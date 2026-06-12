from __future__ import annotations

import tempfile
import unittest
from copy import deepcopy
from pathlib import Path

from docx import Document

from accounting_model import get_account_ledger, get_trace
from document_generator import generar_documento_completo
from financial_model import build_financial_model, result_to_json


def sample_payload():
    return {
        "client": {
            "nombre_completo": "Cliente Prueba",
            "cedula": "001-010101-0000A",
            "estado_civil": "casada",
            "profesion": "Comerciante",
            "sexo": "Femenino",
            "domicilio": "Managua",
            "direccion_negocio": "Managua",
            "banco": "BAC",
            "fecha_certificacion": "2026-05-02",
        },
        "period": {
            "end_month": "2026-04",
            "months": 6,
            "exchange_rate": 36.6243,
            "seed": "modelo-prueba",
        },
        "income": {
            "base_income_usd": 100000,
            "income_variability_pct": 10,
            "cost_pct": 70,
            "cost_variability_pct": 5,
            "cash_sales_pct": 85,
        },
        "movements": {
            "purchase_base_usd": 110000,
            "purchase_variability_pct": 10,
            "events": [
                {"month": "2026-02", "account": "owner_withdrawal", "amount": 250000, "currency": "nio"},
                {"month": "2025-12", "account": "asset_vehicle", "amount": 12000, "currency": "usd"},
            ],
        },
    }


def month_column(df, month):
    for col in df.columns:
        text = col.strftime("%Y-%m") if hasattr(col, "strftime") else str(col)[:7]
        if text == month:
            return col
    raise AssertionError(f"Mes no encontrado: {month}")


def month_columns(df):
    out = []
    for col in df.columns:
        text = col.strftime("%Y-%m") if hasattr(col, "strftime") else str(col)
        if len(text) >= 7 and text[:7].count("-") == 1 and text[:7].replace("-", "").isdigit():
            out.append(text[:7])
    return out


def esf_value(result, label, month):
    col = month_column(result.df_esf_mensual, month)
    row = result.df_esf_mensual[result.df_esf_mensual["Descripcion"] == label].iloc[0]
    return round(float(row[col]))


def full_esf_value(result, label, month):
    col = month_column(result.df_esf_mensual_full, month)
    row = result.df_esf_mensual_full[result.df_esf_mensual_full["Descripcion"] == label].iloc[0]
    return round(float(row[col]))


def account_movement_value(result, account, movement, month):
    col = month_column(result.df_movimiento_cuentas, month)
    rows = result.df_movimiento_cuentas[
        (result.df_movimiento_cuentas["Cuenta"] == account)
        & (result.df_movimiento_cuentas["Movimiento"] == movement)
    ]
    return round(float(rows.iloc[0][col]))


def _docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class FinancialModelTest(unittest.TestCase):
    def test_model_is_reproducible_and_valid(self):
        one = build_financial_model(sample_payload())
        two = build_financial_model(sample_payload())

        self.assertTrue(one.validations["er"]["ok"])
        self.assertTrue(one.validations["esf"]["ok"])
        self.assertTrue(one.validations["balance"]["ok"])
        self.assertEqual(one.summary, two.summary)
        self.assertTrue(one.df_er.equals(two.df_er))
        self.assertTrue(one.df_esf_mensual.equals(two.df_esf_mensual))

    def test_monthly_revenue_override_used_when_present(self):
        payload = sample_payload()
        baseline = build_financial_model(payload)
        payload["income"]["monthly_overrides"] = [{"month": "2026-02", "revenue_usd": 123456}]

        result = build_financial_model(payload)

        col = month_column(result.df_er, "2026-02")
        self.assertEqual(round(float(result.df_er[result.df_er["Descripcion"] == "Ingresos"].iloc[0][col])), round(123456 * 36.6243))
        self.assertEqual(result.summary["exact_revenue_months"], ["2026-02"])
        self.assertEqual(result.metadata["exact_revenue_months"], ["2026-02"])
        self.assertEqual(result.statement_blocks[0]["summary"]["exact_revenue_months"], ["2026-02"])
        baseline_col = month_column(baseline.df_er, "2026-03")
        result_col = month_column(result.df_er, "2026-03")
        self.assertEqual(
            round(float(baseline.df_er[baseline.df_er["Descripcion"] == "Ingresos"].iloc[0][baseline_col])),
            round(float(result.df_er[result.df_er["Descripcion"] == "Ingresos"].iloc[0][result_col])),
        )

    def test_monthly_cogs_override_independent_from_revenue(self):
        payload = sample_payload()
        baseline = build_financial_model(payload)
        payload["income"]["monthly_overrides"] = [{"month": "2026-01", "cogs_usd": 64000}]

        result = build_financial_model(payload)

        col = month_column(result.df_er, "2026-01")
        self.assertEqual(round(float(result.df_er[result.df_er["Descripcion"] == "(-) Costo de ventas"].iloc[0][col])), round(64000 * 36.6243))
        self.assertEqual(result.summary["exact_cogs_months"], ["2026-01"])
        self.assertEqual(result.summary["exact_revenue_months"], [])
        self.assertEqual(
            round(float(baseline.df_er[baseline.df_er["Descripcion"] == "Ingresos"].iloc[0][col])),
            round(float(result.df_er[result.df_er["Descripcion"] == "Ingresos"].iloc[0][col])),
        )

    def test_payload_without_overrides_identical_to_baseline(self):
        payload = sample_payload()
        baseline = build_financial_model(payload)
        with_empty_overrides = deepcopy(payload)
        with_empty_overrides["income"]["monthly_overrides"] = []

        result = build_financial_model(with_empty_overrides)

        self.assertEqual(baseline.summary, result.summary)
        self.assertTrue(baseline.df_er.equals(result.df_er))
        self.assertTrue(baseline.df_esf_mensual.equals(result.df_esf_mensual))
        self.assertTrue(baseline.df_movimientos.equals(result.df_movimientos))

    def test_negative_income_override_ignored_with_warning(self):
        payload = sample_payload()
        baseline = build_financial_model(payload)
        payload["income"]["monthly_overrides"] = [{"month": "2026-01", "revenue_usd": -100, "cogs_usd": -50}]

        result = build_financial_model(payload)

        self.assertEqual(result.summary["exact_revenue_months"], [])
        self.assertEqual(result.summary["exact_cogs_months"], [])
        self.assertTrue(any(w["type"] == "negative_amount" for w in result.metadata["income_override_warnings"]))
        self.assertEqual(
            round(float(baseline.df_er[baseline.df_er["Descripcion"] == "Ingresos"].iloc[0][month_column(baseline.df_er, "2026-01")])),
            round(float(result.df_er[result.df_er["Descripcion"] == "Ingresos"].iloc[0][month_column(result.df_er, "2026-01")])),
        )

    def test_accounting_layer_generates_balanced_opening_and_monthly_vouchers(self):
        result = build_financial_model(sample_payload())
        accounting = result.accounting

        self.assertGreater(accounting["summary"]["voucher_count"], 0)
        opening = accounting["vouchers"][0]
        self.assertEqual(opening["type"], "opening")
        self.assertTrue(opening["balanced"])
        self.assertTrue(all(voucher["balanced"] for voucher in accounting["vouchers"]))
        self.assertTrue(any(voucher["type"] == "sales" for voucher in accounting["vouchers"]))
        self.assertTrue(any(voucher["type"] == "cogs" for voucher in accounting["vouchers"]))
        self.assertTrue(any(voucher["type"] == "expenses" for voucher in accounting["vouchers"]))
        self.assertIn("Resultados Acumulados", accounting["accounts"])
        expense_voucher = next(voucher for voucher in accounting["vouchers"] if voucher["type"] == "expenses")
        expense_accounts = {line["account"] for line in expense_voucher["lines"]}
        self.assertIn("Sueldos y Salarios", expense_accounts)
        self.assertIn("Servicios Publicos", expense_accounts)
        self.assertNotIn("Gastos Operativos", expense_accounts)
        close_voucher = next(voucher for voucher in accounting["vouchers"] if voucher["type"] == "year_close")
        close_accounts = {line["account"] for line in close_voucher["lines"]}
        self.assertIn("Sueldos y Salarios", close_accounts)

    def test_parent_expense_ledger_rolls_up_child_accounts(self):
        result = build_financial_model(sample_payload())

        child_rows = get_account_ledger(result.accounting, "Sueldos y Salarios")
        parent_rows = get_account_ledger(result.accounting, "Sueldos")

        self.assertTrue(child_rows)
        self.assertTrue(parent_rows)
        self.assertEqual(
            sum(row["debit"] for row in child_rows),
            sum(row["debit"] for row in parent_rows),
        )
        trace = get_trace(result.accounting, "Sueldos", result.summary["months"][0])
        self.assertGreater(trace["debits"], 0)

    def test_operating_expense_rollup_equals_sum_of_subcuentas(self):
        """Invariante post voucher-por-subcuenta: el mayor del rubro
        'Gastos Operativos' debe igualar la suma de mayores de todas las
        subcuentas hijas declaradas en EXPENSE_ROLLUP_CHILDREN. Protege
        contra que el motor pierda lineas en el cambio a generacion por
        subcuenta, y garantiza que un periodo abierto despues del cambio
        sigue produciendo el mismo total de ER."""
        from accounting_model import EXPENSE_ROLLUP_CHILDREN

        result = build_financial_model(sample_payload())
        children = EXPENSE_ROLLUP_CHILDREN["Gastos Operativos"]
        sub_total = 0.0
        any_movement = False
        for label in children:
            rows = get_account_ledger(result.accounting, label)
            sub_total += sum(row["debit"] for row in rows)
            if rows:
                any_movement = True
        self.assertTrue(any_movement, "ninguna subcuenta hoja tiene movimientos")

        rollup_rows = get_account_ledger(result.accounting, "Gastos Operativos")
        rollup_total = sum(row["debit"] for row in rollup_rows)
        self.assertAlmostEqual(rollup_total, sub_total, places=0)

    def test_accounting_trace_explains_retained_earnings_year_close(self):
        payload = sample_payload()
        payload["period"] = {
            "start_month": "2025-01",
            "end_month": "2026-04",
            "exchange_rate": 36.6243,
            "seed": "modelo-prueba",
        }
        payload["movements"]["journal_entries"] = [
            {
                "month": "2026-01",
                "debit_account": "current_earnings",
                "credit_account": "retained_earnings",
                "amount": 6_023_510,
                "currency": "nio",
                "entry_type": "year_close_transfer",
                "source": "chat_financiero",
                "instruction_id": "chat_cierre",
                "message": "cierre 2025",
            }
        ]

        result = build_financial_model(payload)
        trace = result.accounting["trace"]["Resultados Acumulados|2026-01"]

        self.assertEqual(trace["opening_balance"], 3_424_750)
        self.assertEqual(trace["credits"], 6_023_510)
        self.assertEqual(trace["closing_balance"], 9_448_260)
        self.assertTrue(any(entry["voucher_id"] for entry in trace["entries"]))
        self.assertTrue(any(v["source"] == "chat_financiero" for v in result.accounting["vouchers"]))

    def test_result_json_exposes_accounting_preview(self):
        result = build_financial_model(sample_payload())
        data = result_to_json(result)

        self.assertIn("accounting", data)
        self.assertGreater(data["accounting"]["summary"]["voucher_count"], 0)
        self.assertIn("ledger", data["accounting"])
        self.assertIn("trace", data["accounting"])

    def test_randomized_rates_stay_inside_configured_ranges(self):
        result = build_financial_model(sample_payload())
        revenue_factors = result.metadata["revenue_factors"]
        cost_rates = result.metadata["cost_rates"]

        self.assertTrue(all(0.90 <= factor <= 1.10 for factor in revenue_factors))
        self.assertTrue(all(0.65 <= rate <= 0.75 for rate in cost_rates))

    def test_cash_flow_preview_reconciles_to_esf_cash(self):
        result = build_financial_model(sample_payload())
        cash_flow_final = result.df_flujo_caja[
            result.df_flujo_caja["Concepto"] == "Saldo final de caja"
        ].iloc[0]
        esf_cash = result.df_esf_mensual[
            result.df_esf_mensual["Descripcion"] == "Efectivo y Equivalentes de Efectivo"
        ].iloc[0]

        for col in result.df_esf_mensual.columns[1:]:
            self.assertEqual(cash_flow_final[col], esf_cash[col])

    def test_account_movement_reconciles_each_account(self):
        result = build_financial_model(sample_payload())
        df = result.df_movimiento_cuentas
        months = list(df.columns[2:])

        for account in df["Cuenta"].unique():
            block = df[df["Cuenta"] == account].set_index("Movimiento")
            for month in months:
                computed = (
                    block.loc["Saldo inicial", month]
                    + block.loc["Aumentos", month]
                    + block.loc["Disminuciones", month]
                )
                self.assertAlmostEqual(computed, block.loc["Saldo final", month], delta=1.0, msg=f"{account} {month}")

    def test_docx_can_be_generated_without_excel(self):
        result = build_financial_model(sample_payload())
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            out_path = Path(tmp.name)
        try:
            generar_documento_completo(
                result.df_esf_mensual,
                result.df_er,
                result.df_datos,
                result.df_certificacion,
                str(out_path),
                incluir_validacion=False,
                esf_tipo="mensual",
            )
            self.assertTrue(out_path.exists())
            self.assertGreater(out_path.stat().st_size, 0)
        finally:
            out_path.unlink(missing_ok=True)

    def test_dynamic_equity_account_flows_to_esf_ledger_and_docx(self):
        payload = sample_payload()
        payload["period"] = {
            "start_month": "2026-01",
            "end_month": "2026-04",
            "exchange_rate": 36.6243,
            "seed": "modelo-prueba",
        }
        payload["accounting"] = {
            "dynamic_accounts": [
                {
                    "code": "reservas_legales",
                    "name": "Reservas Legales",
                    "account_type": "patrimonio",
                    "section": "patrimonio",
                }
            ]
        }
        payload["movements"]["journal_entries"] = [
            {
                "month": "2026-01",
                "debit_account": "capital",
                "credit_account": "Reservas Legales",
                "amount": 250000,
                "currency": "nio",
            }
        ]
        result = build_financial_model(payload)

        self.assertIn("Reservas Legales", set(result.df_esf_mensual_full["Descripcion"]))
        self.assertIn("Reservas Legales", result.accounting["accounts"])
        self.assertTrue(result.validations["balance"]["ok"])

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            out_path = Path(tmp.name)
        try:
            generar_documento_completo(
                result.df_esf_mensual,
                result.df_er,
                result.df_datos,
                result.df_certificacion,
                str(out_path),
                incluir_validacion=False,
                esf_tipo="mensual",
            )
            text = _docx_text(out_path)
            self.assertIn("Reservas Legales", text)
        finally:
            out_path.unlink(missing_ok=True)

    def test_start_end_six_months_uses_single_block(self):
        payload = sample_payload()
        payload["period"] = {
            "start_month": "2025-11",
            "end_month": "2026-04",
            "exchange_rate": 36.6243,
            "seed": "modelo-prueba",
        }

        result = build_financial_model(payload)

        self.assertEqual(result.summary["months"], ["2025-11", "2025-12", "2026-01", "2026-02", "2026-03", "2026-04"])
        self.assertEqual(len(result.statement_blocks), 1)
        self.assertEqual(result.statement_blocks[0]["meta"]["id"], "full_range")
        self.assertEqual(result.statement_blocks[0]["meta"]["label"], "Noviembre 2025-Abril 2026")
        self.assertEqual(month_columns(result.df_er), result.summary["months"])

    def test_more_than_twelve_months_splits_by_calendar_year(self):
        payload = sample_payload()
        payload["period"] = {
            "start_month": "2025-01",
            "end_month": "2026-04",
            "exchange_rate": 36.6243,
            "seed": "modelo-prueba",
        }

        result = build_financial_model(payload)
        metas = [block["meta"] for block in result.statement_blocks]

        self.assertEqual([meta["id"] for meta in metas], ["year_2026", "year_2025"])
        self.assertEqual(metas[0]["label"], "Enero-Abril 2026")
        self.assertEqual(metas[1]["label"], "Enero-Diciembre 2025")
        self.assertEqual(result.summary["months"], ["2026-01", "2026-02", "2026-03", "2026-04"])
        self.assertEqual(result.metadata["full_summary"]["months"][0], "2025-01")
        self.assertEqual(result.metadata["full_summary"]["months"][-1], "2026-04")
        self.assertEqual(len(result.metadata["full_summary"]["months"]), 16)
        self.assertEqual(month_columns(result.df_er), result.summary["months"])
        self.assertTrue(result.validations["balance"]["ok"])

    def test_twenty_eight_months_splits_into_three_blocks(self):
        payload = sample_payload()
        payload["period"] = {
            "start_month": "2024-01",
            "end_month": "2026-04",
            "exchange_rate": 36.6243,
            "seed": "modelo-prueba",
        }

        result = build_financial_model(payload)

        self.assertEqual([block["meta"]["id"] for block in result.statement_blocks], ["year_2026", "year_2025", "year_2024"])
        self.assertEqual(len(result.metadata["full_summary"]["months"]), 28)


    def test_result_json_exposes_block_previews(self):
        payload = sample_payload()
        payload["period"] = {
            "start_month": "2025-01",
            "end_month": "2026-04",
            "exchange_rate": 36.6243,
            "seed": "modelo-prueba",
        }

        data = result_to_json(build_financial_model(payload))

        self.assertEqual([block["id"] for block in data["period_blocks"]], ["year_2026", "year_2025"])
        self.assertIn("full_summary", data)
        self.assertIn("year_2026", data["preview"]["blocks"])
        self.assertIn("year_2025", data["preview"]["blocks"])

    def test_docx_can_be_generated_with_statement_blocks(self):
        payload = sample_payload()
        payload["period"] = {
            "start_month": "2025-01",
            "end_month": "2026-04",
            "exchange_rate": 36.6243,
            "seed": "modelo-prueba",
        }
        result = build_financial_model(payload)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            out_path = Path(tmp.name)
        try:
            generar_documento_completo(
                result.df_esf_mensual,
                result.df_er,
                result.df_datos,
                result.df_certificacion,
                str(out_path),
                incluir_validacion=False,
                esf_tipo="mensual",
                statement_blocks=result.statement_blocks,
            )
            self.assertTrue(out_path.exists())
            self.assertGreater(out_path.stat().st_size, 0)
        finally:
            out_path.unlink(missing_ok=True)























class CapitalInvariantTest(unittest.TestCase):
    """F1-T1: el capital residual del ESF debe coincidir con el transaccional."""

    def test_sane_model_passes_capital_check(self):
        result = build_financial_model(sample_payload())

        self.assertTrue(result.validations["capital"]["ok"])
        self.assertEqual(result.validations["capital"]["errors"], [])

    def test_contributions_reclassification_and_journal_entries_are_tracked(self):
        payload = sample_payload()
        payload["movements"]["events"].extend([
            {"month": "2026-01", "account": "capital_contribution", "amount": 100000, "currency": "nio"},
            {"month": "2026-03", "account": "reclasificacion_capital", "amount": 80000, "currency": "nio"},
        ])
        payload["movements"]["journal_entries"] = [
            {
                "month": "2026-02",
                "debit_account": "capital",
                "credit_account": "cash",
                "amount_nio": 50000,
                "description": "Retiro de capital via asiento del chat",
            }
        ]

        result = build_financial_model(payload)

        self.assertTrue(
            result.validations["capital"]["ok"],
            result.validations["capital"]["errors"],
        )
        self.assertTrue(result.validations["balance"]["ok"])

    def test_capital_check_detects_hidden_imbalance(self):
        # credit_card_new no tiene contrapartida en el modelo: el pasivo sube
        # sin que entre ningun activo ni se registre gasto. Es un descuadre
        # real que el balance_check tradicional no ve (el capital residual lo
        # absorbe) pero que el invariante transaccional reporta con mes y
        # monto exactos.
        payload = sample_payload()
        payload["movements"]["events"].append(
            {"month": "2026-02", "account": "tarjeta", "amount": 50_000, "currency": "nio"}
        )

        result = build_financial_model(payload)
        capital_validation = result.validations["capital"]

        self.assertFalse(capital_validation["ok"])
        months = [error["month"] for error in capital_validation["errors"]]
        self.assertIn("2026-02", months)
        first = next(e for e in capital_validation["errors"] if e["month"] == "2026-02")
        self.assertAlmostEqual(first["difference"], -50_000, delta=2)
        # El descuadre persiste hasta el final del periodo.
        self.assertIn("2026-04", months)
        # Y el balance_check tradicional NO lo detecta (capital residual lo
        # absorbe): esa es exactamente la razon de ser de este invariante.
        self.assertTrue(result.validations["balance"]["ok"])


class OverpaymentTest(unittest.TestCase):
    """F1-T2: pagos por encima del saldo del pasivo se recortan con warning."""

    def test_overpayment_is_capped_with_warning(self):
        baseline = build_financial_model(sample_payload())
        payload = sample_payload()
        payload["movements"]["events"].extend([
            {"month": "2026-02", "account": "abono_tarjeta", "amount": 1_000_000, "currency": "nio"},
            {"month": "2026-03", "account": "abono_personal", "amount": 100_000, "currency": "nio"},
        ])

        result = build_financial_model(payload)

        overpayments = result.validations["overpayments"]
        self.assertFalse(overpayments["ok"])
        cc = next(w for w in overpayments["warnings"] if w["account"] == "credit_cards")
        self.assertEqual(cc["month"], "2026-02")
        self.assertEqual(round(cc["requested"]), 1_000_000)
        self.assertEqual(round(cc["applied"]), 183_122)
        loan = next(w for w in overpayments["warnings"] if w["account"] == "loans_personal")
        self.assertEqual(loan["month"], "2026-03")
        self.assertEqual(round(loan["applied"]), 47_612)

        # El pasivo queda en cero y la caja baja solo por el monto aplicado.
        self.assertEqual(esf_value(result, "Tarjetas de Credito", "2026-02"), 0)
        cash_delta = (
            esf_value(result, "Efectivo y Equivalentes de Efectivo", "2026-02")
            - esf_value(baseline, "Efectivo y Equivalentes de Efectivo", "2026-02")
        )
        self.assertEqual(cash_delta, -183_122)

        # Sin descuadres: todos los invariantes quedan en verde.
        self.assertTrue(result.validations["balance"]["ok"])
        self.assertTrue(result.validations["capital"]["ok"], result.validations["capital"]["errors"])

    def test_exact_payment_produces_no_warning(self):
        payload = sample_payload()
        payload["movements"]["events"].append(
            {"month": "2026-02", "account": "abono_tarjeta", "amount": 183_122, "currency": "nio"}
        )

        result = build_financial_model(payload)

        self.assertTrue(result.validations["overpayments"]["ok"])
        self.assertEqual(esf_value(result, "Tarjetas de Credito", "2026-02"), 0)
        self.assertTrue(result.validations["capital"]["ok"])


if __name__ == "__main__":
    unittest.main()
