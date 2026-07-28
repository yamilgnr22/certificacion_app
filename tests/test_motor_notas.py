"""Tests de las notas integradoras (motor/notas.py + generators/notas.py).

Regla de oro: cada nota cuadra EXACTO contra el ESF de corte.
Casos: Gloria (Tipo A real) y demo Tipo B. Ademas un e2e que genera el DOCX
con notas y verifica que la seccion exista.
"""

from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from motor import certificar_tipo_a, certificar_tipo_b
from motor.notas import construir_notas
from test_motor_golden_gloria import _inputs_gloria
from test_motor_tipo_b import _inputs as _inputs_tipo_b


class NotasGloriaTest(unittest.TestCase):
    def setUp(self):
        self.m = certificar_tipo_a(_inputs_gloria())
        self.notas = construir_notas(self.m)
        self.corte = self.m.esf.corte()

    def test_hay_4_notas(self):
        self.assertEqual([n.numero for n in self.notas], [1, 2, 3, 4])

    def test_nota1_efectivo_cuadra(self):
        self.assertAlmostEqual(self.notas[0].total[-1], self.corte.efectivo, places=2)
        self.assertAlmostEqual(self.notas[0].total[-1], 841_220.0, places=2)

    def test_nota2_inventarios_cuadra(self):
        self.assertAlmostEqual(self.notas[1].total[-1], self.corte.inventarios, places=2)

    def test_nota3_ppe_cuadra(self):
        _, costo, depr, valor = self.notas[2].total
        self.assertAlmostEqual(costo, 1_062_105.0, places=2)  # 366,243 + 695,862
        self.assertAlmostEqual(depr, -54_018.0, places=2)
        self.assertAlmostEqual(valor, 1_008_087.0, places=2)

    def test_nota3_depreciacion_prorrateada_suma_exacta(self):
        suma_depr = round(sum(f[2] for f in self.notas[2].filas), 2)
        self.assertAlmostEqual(suma_depr, self.corte.depreciacion_acumulada, places=2)
        # valor en libros por fila = costo + depr
        for f in self.notas[2].filas:
            self.assertAlmostEqual(f[3], round(f[1] + f[2], 2), places=2)

    def test_nota4_pasivos_cuadra(self):
        self.assertAlmostEqual(self.notas[3].total[-1], self.corte.total_pasivos, places=2)
        self.assertAlmostEqual(self.notas[3].total[-1], 60_430.0, places=2)
        # Solo cuentas con saldo (Gloria: solo tarjetas)
        self.assertEqual(len(self.notas[3].filas), 1)
        self.assertEqual(self.notas[3].filas[0][0], "Tarjetas de Credito")

    def test_fecha_corte_en_totales(self):
        self.assertIn("31/05/2026", self.notas[0].total[0])


class NotasTipoBTest(unittest.TestCase):
    def setUp(self):
        self.m = certificar_tipo_b(_inputs_tipo_b(con_inventario=True))
        self.notas = construir_notas(self.m)
        self.corte = self.m.esf.corte()

    def test_totales_cuadran_contra_corte(self):
        self.assertAlmostEqual(self.notas[0].total[-1], self.corte.efectivo, places=2)
        self.assertAlmostEqual(self.notas[1].total[-1], self.corte.inventarios, places=2)
        self.assertAlmostEqual(self.notas[2].total[-1], round(
            self.corte.bienes_inmuebles + self.corte.mobiliario_equipos
            + self.corte.vehiculos + self.corte.depreciacion_acumulada, 2
        ), places=2)
        self.assertAlmostEqual(self.notas[3].total[-1], self.corte.total_pasivos, places=2)

    def test_sin_pasivos_muestra_linea_cero(self):
        self.assertEqual(self.notas[3].filas, [["Sin pasivos al corte", 0.0]])


class NotasCuentasBancariasTest(unittest.TestCase):
    """Nota 1 desglosada: cuentas bancarias + 'Efectivo en Caja' residuo.

    Regla del CPA: caja = efectivo ESF - suma de cuentas (determinista);
    negativa = BLOQUEANTE. USD se convierte con la tasa del periodo."""

    def setUp(self):
        self.m = certificar_tipo_a(_inputs_gloria())  # efectivo corte = 841,220
        self.tc = self.m.inputs.periodo.tasa_cambio

    def test_desglose_con_caja_residual(self):
        cuentas = [
            {"banco": "LAFISE", "tipo": "Cuenta de Ahorro", "moneda": "NIO",
             "numero": "106012140", "saldo": 500_000},
            {"banco": "BAC", "tipo": "Cuenta Corriente", "moneda": "USD",
             "numero": "367157328", "saldo": 5_000},
        ]
        notas = construir_notas(self.m, cuentas_bancarias=cuentas)
        n1 = notas[0]
        usd_nio = round(5_000 * self.tc)
        # Fila 1 = Efectivo en Caja (residuo), luego las cuentas
        self.assertEqual(n1.filas[0][0], "Efectivo en Caja")
        self.assertEqual(n1.filas[0][1], 841_220 - 500_000 - usd_nio)
        self.assertEqual(n1.filas[1][0], "LAFISE Cuenta de Ahorro NIO No. 106012140")
        self.assertEqual(n1.filas[1][1], 500_000)
        self.assertEqual(n1.filas[2][0], "BAC Cuenta Corriente USD No. 367157328")
        self.assertEqual(n1.filas[2][1], usd_nio)
        # La suma de filas = total = efectivo del ESF, EXACTO
        self.assertEqual(sum(f[1] for f in n1.filas), 841_220)
        self.assertEqual(n1.total[-1], 841_220)

    def test_cuentas_igual_al_efectivo_omite_caja(self):
        cuentas = [{"banco": "LAFISE", "tipo": "Cuenta de Ahorro", "moneda": "NIO",
                    "numero": "1", "saldo": 841_220}]
        n1 = construir_notas(self.m, cuentas_bancarias=cuentas)[0]
        self.assertEqual(len(n1.filas), 1)  # sin fila de caja (residuo 0)
        self.assertEqual(sum(f[1] for f in n1.filas), 841_220)

    def test_caja_negativa_es_bloqueante(self):
        from motor.notas import NotasError

        cuentas = [{"banco": "LAFISE", "tipo": "Cuenta de Ahorro", "moneda": "NIO",
                    "numero": "1", "saldo": 900_000}]  # > 841,220
        with self.assertRaises(NotasError) as ctx:
            construir_notas(self.m, cuentas_bancarias=cuentas)
        self.assertIn("negativo", str(ctx.exception))

    def test_sin_cuentas_linea_unica_de_siempre(self):
        n1 = construir_notas(self.m)[0]
        self.assertEqual(n1.filas, [["Efectivo en Caja y Bancos", 841_220.0]])


class NotasDocxE2ETest(unittest.TestCase):
    def test_docx_incluye_seccion_de_notas(self):
        from docx import Document
        from document_generator import generar_documento_completo

        m = certificar_tipo_a(_inputs_gloria())
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            out = tmp.name
        try:
            generar_documento_completo(
                m.df_esf_corte, m.df_er, m.df_datos, m.df_certificacion, out,
                incluir_validacion=False, esf_tipo="corte",
                notas_data=construir_notas(m),
            )
            doc = Document(out)
            textos = [p.text for p in doc.paragraphs]
            self.assertTrue(any("INTEGRACION DEL EFECTIVO" in t for t in textos))
            self.assertTrue(any("INTEGRACION DE PASIVOS" in t for t in textos))
            # La tabla de la nota 1 trae el total cuadrado
            montos = {c.text for tb in doc.tables for r in tb.rows for c in r.cells}
            self.assertIn("841,220", montos)
        finally:
            Path(out).unlink(missing_ok=True)


if __name__ == "__main__":
    unittest.main()
