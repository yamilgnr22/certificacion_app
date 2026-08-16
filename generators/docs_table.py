# generators/docs_table.py
"""
Secciones de imágenes del DOCX: “Documentos del cliente” (siempre va) y
“Fotografías del Negocio” (opcional).

Dos modos por sección:
  • CON imágenes (motor V2): grilla de 2 por fila (8.5 cm de ancho cada una).
    Cada ítem puede ser una ruta (str) o un dict {tipo, path}; con tipos, las
    filas se emparejan semánticamente: cédula front|back, luego cada
    matricula_N junto a su soporte_N, y el resto de a dos.
  • SIN imágenes: tabla vacía 3×5 como siempre (respaldo para pegar a mano).
"""

import io
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from typing import List, Optional, Tuple

from word_helpers import apply_paragraph_style, set_row_height, set_cell_border


def _titulo(doc: Document, texto: str = "Documentos del cliente") -> None:
    doc.add_page_break()
    titulo = doc.add_paragraph(texto)
    apply_paragraph_style(
        titulo,
        font_name="Arial",
        font_size=12,
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        line_spacing=1.15,
    )
    doc.add_paragraph()  # pequeño espacio


def _tabla_vacia(doc: Document) -> None:
    table = doc.add_table(rows=5, cols=3)
    table.autofit = False
    col_w = [8.5, 1.0, 8.5]  # cm
    for j, w in enumerate(col_w):
        for row in table.rows:
            row.cells[j].width = Cm(w)
    row_h = [5.0, 1.0, 5.0, 1.0, 5.0]  # cm
    for row, h in zip(table.rows, row_h):
        set_row_height(row, int(h * 1440 / 2.54))
    no_border = {"sz": "0", "val": "nil"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=no_border, bottom=no_border, left=no_border, right=no_border)


def _normalizar(items) -> List[Tuple[str, str]]:
    """Lista uniforme [(tipo, path)] a partir de rutas sueltas o dicts."""
    out: List[Tuple[str, str]] = []
    for it in items or []:
        if isinstance(it, dict):
            out.append((str(it.get("tipo") or "otro"), str(it.get("path") or "")))
        elif isinstance(it, (tuple, list)) and len(it) == 2:
            out.append((str(it[0]), str(it[1])))
        else:
            out.append(("otro", str(it)))
    return out


def _filas_pareadas(items) -> List[Tuple[Optional[str], Optional[str]]]:
    """Filas (izquierda, derecha) con el layout de la certificación:

    fila 1: cédula frente | cédula reverso
    fila N: matricula_N   | soporte_N  (constancia de trámite)
    resto:  de a dos en el orden recibido (matricula legado, otros...).
    None = celda vacía (p. ej. matrícula sin soporte aún)."""
    disponibles = _normalizar(items)

    def tomar(tipo: str) -> Optional[str]:
        i = next((i for i, (t, _) in enumerate(disponibles) if t == tipo), None)
        return disponibles.pop(i)[1] if i is not None else None

    filas: List[Tuple[Optional[str], Optional[str]]] = []
    izq, der = tomar("cedula_front"), tomar("cedula_back")
    if izq or der:
        filas.append((izq, der))
    for n in ("1", "2", "3"):
        m, s = tomar(f"matricula_{n}"), tomar(f"soporte_{n}")
        if m or s:
            filas.append((m, s))
    resto = [p for _, p in disponibles]
    for i in range(0, len(resto), 2):
        filas.append((resto[i], resto[i + 1] if i + 1 < len(resto) else None))
    return filas


def _abrir_enderezada(ruta: str):
    """Devuelve lo que hay que insertar en el DOCX, ya derecho.

    Las fotos de celular guardan la rotacion en un tag EXIF en vez de rotar
    los pixeles: el visor de Windows la aplica y se ven bien, pero Word NO,
    asi que la cedula entraba acostada y, escalada por ancho, ocupaba media
    pagina. Aca se aplica la rotacion y se entrega la imagen en memoria; el
    archivo original del cliente no se toca.

    Sin EXIF de orientacion (o si Pillow no puede leerla) devuelve la ruta
    tal cual, que es el camino de siempre.
    """
    try:
        from PIL import Image, ImageOps

        with Image.open(ruta) as im:
            if (im.getexif() or {}).get(274, 1) in (1, None):
                return ruta  # ya esta derecha: se inserta el archivo original
            derecha = ImageOps.exif_transpose(im)
            buffer = io.BytesIO()
            formato = im.format or "PNG"
            if formato.upper() in {"JPEG", "JPG"} and derecha.mode not in ("RGB", "L"):
                derecha = derecha.convert("RGB")
            derecha.save(buffer, format=formato)
            buffer.seek(0)
            return buffer
    except Exception:
        return ruta


def _grilla_imagenes(doc: Document, imagenes) -> None:
    """Incrusta las imágenes 2 por fila (columna separadora de 1 cm al medio)."""
    no_border = {"sz": "0", "val": "nil"}
    for izq, der in _filas_pareadas(imagenes):
        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        for j, w in enumerate([8.5, 1.0, 8.5]):
            table.rows[0].cells[j].width = Cm(w)
        for col, ruta in ((0, izq), (2, der)):
            if ruta is None:
                continue
            parrafo = table.rows[0].cells[col].paragraphs[0]
            parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            try:
                if os.path.isfile(ruta):
                    parrafo.add_run().add_picture(_abrir_enderezada(ruta), width=Cm(8.3))
                else:
                    parrafo.add_run("(imagen no encontrada)")
            except Exception:
                parrafo.add_run("(imagen ilegible)")
        for cell in table.rows[0].cells:
            set_cell_border(cell, top=no_border, bottom=no_border, left=no_border, right=no_border)
        doc.add_paragraph()  # aire entre filas de imágenes


def generar_tabla_docs_cliente(doc: Document, imagenes: Optional[list] = None) -> None:
    _titulo(doc)
    if imagenes:
        _grilla_imagenes(doc, imagenes)
    else:
        _tabla_vacia(doc)
    doc.add_page_break()


def generar_fotos_negocio(doc: Document, imagenes: Optional[list] = None) -> None:
    """Hoja opcional “Fotografías del Negocio” (fotos del local del cliente).

    Con imágenes las incrusta de a dos; sin imágenes deja la tabla vacía
    para pegado manual en Word (mismo respaldo que Documentos del cliente)."""
    _titulo(doc, "Fotografías del Negocio")
    if imagenes:
        _grilla_imagenes(doc, imagenes)
    else:
        _tabla_vacia(doc)
    doc.add_page_break()