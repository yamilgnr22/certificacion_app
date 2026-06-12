# certificacion.py
from datetime import datetime
import calendar
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from num2words import num2words

from config_cpa import load_cpa_profile
from word_helpers import (
    apply_paragraph_style,
    add_paragraph_border,
    add_page_number,
)
from .utils import extract_cert_fields

def _fecha_a_palabras(fecha):
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    dia = num2words(fecha.day, lang="es")
    mes = meses[fecha.month - 1]
    anio = num2words(fecha.year, lang="es")
    return f"{dia} días del mes de {mes} del año {anio}"

def _periodo_certificacion(inicio, fin):
    meses_es = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    last_day_fin = calendar.monthrange(fin.year, fin.month)[1]
    return (
        f"Para el periodo comprendido del 1ro de {meses_es[inicio.month - 1]} "
        f"del año {inicio.year} al {last_day_fin} de {meses_es[fin.month - 1]} "
        f"del año {fin.year}"
    )

def generar_certificacion(doc: Document, dfc: pd.DataFrame, profile=None):
    """
    Inserta en `doc` la sección de certificación completa:
    encabezado, cuerpo y pie de página.
    """
    # ————————————————————— Datos básicos (robustos) —————————————————————
    cpa = profile or load_cpa_profile()
    cert = extract_cert_fields(dfc)
    nombre_cliente = cert.get("nombre")
    apellido_cliente = cert.get("apellido")
    nombre_completo = (cert.get("nombre_completo") or f"{nombre_cliente or ''} {apellido_cliente or ''}").strip()
    cedula = cert.get("cedula")
    inicio = cert.get("inicio")
    fin    = cert.get("fin")
    estado_civil = cert.get("estado_civil")
    profesion    = cert.get("profesion")
    sexo         = cert.get("sexo")
    domicilio    = cert.get("domicilio")
    banco        = cert.get("banco")
    fecha_cert   = cert.get("fecha_certificacion")

    ingresos_brutos   = cert.get("ingresos_brutos")
    ingresos_promedio = cert.get("ingresos_promedio")
    utilidad_periodo  = cert.get("utilidad_periodo")
    utilidad_promedio = cert.get("utilidad_promedio")

    ingresos_brutos_palabras   = num2words(ingresos_brutos,   lang="es")
    ingresos_promedio_palabras = num2words(ingresos_promedio, lang="es")
    utilidad_periodo_palabras  = num2words(utilidad_periodo,  lang="es")
    utilidad_promedio_palabras = num2words(utilidad_promedio, lang="es")
    if fecha_cert is None:
        raise ValueError("Fecha de certificación no válida en hoja 'Certificacion'")
    fecha_certificacion_palabras = _fecha_a_palabras(fecha_cert)

    genero_cliente   = "la señora" if str(sexo).strip().lower() == "femenino" else "el señor"
    genero_cliente_2 = "de la señora" if str(sexo).strip().lower() == "femenino" else "del señor"

    ingresos_brutos_fmt   = f"{ingresos_brutos:,.0f}"
    ingresos_promedio_fmt = f"{ingresos_promedio:,.0f}"
    utilidad_periodo_fmt  = f"{utilidad_periodo:,.0f}"
    utilidad_promedio_fmt = f"{utilidad_promedio:,.0f}"

    if inicio is None or fin is None:
        raise ValueError("Fechas de inicio/fin no válidas en hoja 'Certificacion'")
    periodo_certificacion = _periodo_certificacion(inicio, fin)

    # ————————————————————— Encabezado y pie de página —————————————————————
    for section in doc.sections:
        # Márgenes
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

        # — Encabezado —
        header = section.header
        if header.paragraphs:
            p_h = header.paragraphs[0]
            p_h.clear()
        else:
            p_h = header.add_paragraph()
        run = p_h.add_run(f"{cpa.titulo_corto} {cpa.nombre}\n")
        run.bold = True
        run.font.name = "Abadi"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(20, 47, 80)
        run2 = p_h.add_run(f"Contador Público Autorizado No. {cpa.numero_cpa}")
        run2.font.name = "Abadi Extra Light"
        run2.font.size = Pt(9)
        apply_paragraph_style(p_h, font_name="Arial", font_size=9)
        add_paragraph_border(p_h, "bottom")
        p_h.paragraph_format.space_after = Pt(12)

        # — Pie de página —
        footer = section.footer
        if footer.paragraphs:
            p_f = footer.paragraphs[0]
            p_f.clear()
        else:
            p_f = footer.add_paragraph()
        p_f.add_run(f"📞 {cpa.telefono}   📧 {cpa.email}\t\tPágina ")
        apply_paragraph_style(p_f, font_name="Abadi Extra Light", font_size=8)
        add_paragraph_border(p_f, "top")
        add_page_number(p_f)
        for run in p_f.runs:
            run.font.name = "Abadi Extra Light"
            run.font.size = Pt(8)

    # ————————————————————— Cuerpo de certificación —————————————————————
    # Título centrado
    p_title = doc.add_paragraph("CERTIFICACIÓN DEL CONTADOR PÚBLICO INDEPENDIENTE")
    apply_paragraph_style(
        p_title,
        font_name="Arial",
        font_size=8,
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        line_spacing=1,
    )
    p_title.paragraph_format.space_after = Pt(12)

    # Texto principal
    cuerpo = [
            f"Yo, {cpa.nombre}, mayor de edad, {cpa.estado_civil}, {cpa.titulo}, del domicilio de {cpa.domicilio}, identificado con cédula de identidad número {cpa.cedula}, en mi calidad de Contador Público Autorizado para ejercer la profesión por el excelentísimo Ministerio de Educación Cultura y Deporte, bajo acuerdo {cpa.acuerdo_cpa}, fechado {cpa.fecha_acuerdo}, por el quinquenio que finalizará el {cpa.fin_quinquenio}; expreso a través de este documento, que conforme a las leyes vigentes del país, en mi carácter de profesional de Contaduría Pública y en representación propia, que fui contratado para Certificar el Estado de Situación Financiera y el Estado de Resultados {periodo_certificacion}, preparados por {genero_cliente} {nombre_completo}, mayor de edad, {estado_civil}, {profesion}, quien se identifica con la cédula de identidad No. {cedula}, con domicilio en {domicilio}.",
            f"Responsabilidad: {genero_cliente} {nombre_completo}, es responsable sobre la información suministrada y reflejada en los estados financieros presentados para la Certificación, mi responsabilidad consiste en Certificar que las cifras contenidas en dichos estados financieros están conformes con los registros contables llevados {genero_cliente} {nombre_completo}.",
            f"Mi trabajo fue realizado de acuerdo con la Normativa sobre “Trabajo previamente convenidos”, y los objetivos que se persiguieron con dicho trabajo fueron los siguientes:",
            # Estos tres párrafos se mostrarán con viñeta
            f"•\tDeterminar que los estados financieros presentados por {genero_cliente} {nombre_completo} fueron preparados de acuerdo con los registros contables llevados para registrar las operaciones.",
            f"•\tDeterminar que los registros contables contenidos en los estados financieros {genero_cliente_2} {nombre_completo} se encuentran de acuerdo con principios de contabilidad generalmente aceptados en Nicaragua.",
            f"•\tDeterminar si los ingresos netos derivados de la actividad económica {genero_cliente_2} {nombre_completo} se encuentran presentados de forma razonable de conformidad con sus registros contables.",
            f"Para lograr los objetivos efectué una revisión selectiva de los registros contables a fin de determinar que estos estaban efectuados de acuerdo con Principios de Contabilidad Generalmente Aceptados en Nicaragua.",
            f"Mi trabajo proporciona una base razonable para Certificar que las cifras contenidas en los Estados Financieros {genero_cliente_2} {nombre_completo}, han sido preparados de acuerdo con los registros contables de sus operaciones a la fecha anteriormente indicada. En este sentido, sobre la base del trabajo que efectué, Certifico que:",
            # Los siguientes dos párrafos deben tener viñeta también:
            f"•\tLos ingresos brutos para el periodo revisado ascendieron a NIO {ingresos_brutos_fmt} ({ingresos_brutos_palabras} córdobas), que da como resultado un promedio mensual de ingresos brutos de NIO {ingresos_promedio_fmt} ({ingresos_promedio_palabras} córdobas).",
            f"•\tLas utilidades netas del periodo revisado (después de deducir costos y gastos) fueron de NIO {utilidad_periodo_fmt} ({utilidad_periodo_palabras} córdobas), que da como resultado un promedio mensual de utilidades netas de NIO {utilidad_promedio_fmt} ({utilidad_promedio_palabras} córdobas).",
            f"Se adjuntan a esta Certificación, el Estado de Resultados, el Estado de Situación Financiera y los anexos a los estados financieros, los cuales han sido rubricado y sellado por el suscrito Contador Público Autorizado.",
            f"Esta certificación ha sido solicitada para completar los requisitos bancarios con {banco}, por lo que no debe ser utilizada para otro trámite legal ante cualquier otra institución pública o privada.",
            f"Dado en la ciudad de {cpa.ciudad_emision}, a los {fecha_certificacion_palabras}."
    ]
    bullets = {3, 4, 5, 8, 9}

    for idx, texto in enumerate(cuerpo):
        if idx in bullets:
            # Asegurar que el texto tenga el bullet como carácter, y no estilo de lista
            t = texto.strip()
            if not t.startswith("•"):
                t = "• " + t

            p = doc.add_paragraph(t)

            # Aplicar el estilo de párrafo a 7pt (incluye viñeta)
            apply_paragraph_style(
                p,
                font_name="Arial",
                font_size=7,  # tamaño requerido
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                line_spacing=1.5,
            )

            # Sangría francesa para simular lista
            fmt = p.paragraph_format
            fmt.left_indent = Cm(1.0)        # margen izquierdo del bloque
            fmt.first_line_indent = Cm(-0.5) # cuelga la primera línea con la viñeta

            # Forzar tamaño de fuente = 7 pt en todos los runs (incluida la "•")
            for run in p.runs:
                run.font.size = Pt(7)
                run.font.name = "Arial"

        else:
            p = doc.add_paragraph(texto)
            apply_paragraph_style(
                p,
                font_name="Arial",
                font_size=7,
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                line_spacing=1.5,
            )

        # Espaciado entre párrafos
        p.paragraph_format.space_after = Pt(12)

    for _ in range(2):
            p_blank = doc.add_paragraph("")
            # ► interlineado
            p_blank.paragraph_format.line_spacing = 1.5
            # ► sin espacio adicional
            p_blank.paragraph_format.space_after  = Pt(0)


    # ————————————————————— Firmas y salto de página —————————————————————
    for line in [
        cpa.nombre,
        f"Cédula de identidad {cpa.cedula}",
        f"Contador Público Autorizado Nº {cpa.numero_cpa}",
    ]:
        p = doc.add_paragraph(line)
        apply_paragraph_style(
            p,
            font_name="Arial",
            font_size=7,
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        p.paragraph_format.space_after = Pt(0)

    doc.add_page_break()
