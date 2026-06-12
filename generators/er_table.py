# generators/er_table.py
"""
Sección del Estado de Resultados con encabezados “Ene‑25, Feb‑25, …”
y bloque de firmas alineado con tabulaciones.
"""

import calendar
from datetime import datetime
import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt
from config_cpa import load_cpa_profile
from word_helpers import (
    apply_paragraph_style,
    set_vertical_alignment,
    set_row_height,
    set_cell_border,
)
from .utils import extract_cert_fields

# ───────────────────────────── utilidades fecha ──────────────────────────────
SPANISH_MONTHS = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]
SPANISH_ABBR = {
    1: "Ene", 2: "Feb", 3: "Mar", 4: "Abr",
    5: "May", 6: "Jun", 7: "Jul", 8: "Ago",
    9: "Sep", 10: "Oct", 11: "Nov", 12: "Dic",
}


def _periodo(inicio: datetime, fin: datetime) -> str:
    last_day = calendar.monthrange(fin.year, fin.month)[1]
    return (
        f"Para el periodo comprendido del 1ro de {SPANISH_MONTHS[inicio.month-1]} "
        f"del año {inicio.year} al {last_day} de {SPANISH_MONTHS[fin.month-1]} "
        f"del año {fin.year}"
    )


def _hdr(col) -> str:
    """Devuelve 'Ene‑25' si la cabecera es fecha; si no, str(col)."""
    try:
        dt = pd.to_datetime(col, errors="raise")
        return f"{SPANISH_ABBR[dt.month]}-{str(dt.year)[-2:]}"
    except Exception:
        return str(col)


# ───────────────────────── función principal ─────────────────────────────────
def generar_tabla_er(doc, df_er: pd.DataFrame, df_cert: pd.DataFrame) -> None:
    # 1) Datos de cabecera (robustos)
    cert = extract_cert_fields(df_cert)
    nombre   = cert.get("nombre")
    apellido = cert.get("apellido")
    nombre_completo = (cert.get("nombre_completo") or f"{nombre or ''} {apellido or ''}").strip()
    cedula   = cert.get("cedula")
    direccion = cert.get("direccion_negocio")
    ini = cert.get("inicio")
    fin = cert.get("fin")
    periodo_txt = _periodo(ini, fin) if (ini is not None and fin is not None) else ""

    # 2) Encabezados centrados
    for txt, size, bold in [
        (f"{nombre_completo}", 12, True),
        (str(cedula), 12, True),
        (str(direccion), 8, False),
        ("Estado de Resultados", 8, False),
        ("Expresado en córdobas", 8, False),
        (periodo_txt, 8, False),
    ]:
        p = doc.add_paragraph(txt)
        apply_paragraph_style(
            p, font_name="Arial", font_size=size, bold=bold,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, line_spacing=1
        )
    doc.add_paragraph()                     # espacio antes de la tabla

    # 3) Preparar DataFrame
    df = (
        df_er.copy()
             .drop(index=[0, 1, 2, 3, 4])
             .reset_index(drop=True)
             .drop(columns=[df_er.columns[1]])
             .fillna("")
    )
    for col in df.columns[1:]:
        df[col] = df[col].apply(
            lambda x: f"{x:,.0f}" if isinstance(x, (int, float)) else x
        )

    # 4) Crear tabla
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 4a) Encabezados columna
    for j, col in enumerate(df.columns):
        c = table.rows[0].cells[j]
        c.text = _hdr(col)
        apply_paragraph_style(
            c.paragraphs[0], font_name="Arial", font_size=7, bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
            space_before=0, space_after=0, line_spacing=1
        )
        set_vertical_alignment(c, "center")
        set_cell_border(
            c,
            top={"sz": "3", "val": "single", "color": "000000"},
            bottom={"sz": "3", "val": "single", "color": "000000"},
        )

    # 4b) Filas de datos
    no_indent = ["Descripción", "(=) Ingresos Brutos", "Total gastos operativos"]
    in_gastos_detalle = False                                       # ← bandera sangría
    for i, row in df.iterrows():
        key = str(row.iloc[0]).strip()

        # actualizar la bandera de detalle (se aplica a la *siguiente* fila)
        if key == "(-) Gastos operativos":
            in_gastos_detalle = True
        elif key == "Total gastos operativos":
            in_gastos_detalle = False

        cells = table.rows[i + 1].cells
        for j, val in enumerate(row):
            cell = cells[j]
            cell.text = str(val)

            # --- estilo base (negrita & alineación) --------------------------
            if key in no_indent:
                align, bold = (WD_PARAGRAPH_ALIGNMENT.LEFT, True) if j == 0 else (
                    WD_PARAGRAPH_ALIGNMENT.RIGHT, True
                )
            else:
                align, bold = (WD_PARAGRAPH_ALIGNMENT.LEFT, False) if j == 0 else (
                    WD_PARAGRAPH_ALIGNMENT.RIGHT, False
                )

            # --- sangría suave para detalle de gastos operativos ------------
            indent_cm = (
                0.15
                if in_gastos_detalle
                and j == 0
                and key not in ["(-) Gastos operativos", "Total gastos operativos"]
                else 0
            )

            apply_paragraph_style(
                cell.paragraphs[0],
                font_name="Arial", font_size=7, bold=bold,
                alignment=align, space_before=0, space_after=0,
                line_spacing=1, indent=indent_cm
            )
            set_vertical_alignment(cell, "center")
        set_row_height(table.rows[i + 1], 283)                     # ~0.5 cm

    # 5) Bordes especiales
    for row in table.rows[1:]:
        first = row.cells[0].text.strip()
        if first in ["(=) Ingresos Brutos", "Total gastos operativos"]:
            for cell in row.cells:
                set_cell_border(cell, top={"sz": "3", "val": "single", "color": "000000"})
    for cell in table.rows[n_rows].cells:                          # borde doble final
        set_cell_border(
            cell,
            top={"sz": "3", "val": "single", "color": "000000"},
            bottom={"sz": "3", "val": "double", "color": "000000"},
        )

    # 5b) Negrita en toda la fila “Ingresos/Utilidad Neta”
    for row in table.rows[1:]:
        if row.cells[0].text.strip().lower().startswith("ingresos/utilidad"):
            for cell in row.cells:
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            break

    # 6) Anchos
    for r in table.rows:
        for k, c in enumerate(r.cells):
            c.width = Cm(4) if k == 0 else Cm(1.8)
    set_row_height(table.rows[0], 397)

    # ───────────── 4 PÁRRAFOS VACÍOS (interlineado 1.5) ──────────────
    for _ in range(4):
        p_blank = doc.add_paragraph("")
        p_blank.paragraph_format.line_spacing = 1.5
        p_blank.paragraph_format.space_after = Pt(0)

    # 7) Bloque de firmas
    doc.add_paragraph()
    cpa = load_cpa_profile()
    nombre_firma = nombre_completo or (f"{nombre or ''} {apellido or ''}".strip()) or ""
    firma = doc.add_paragraph(
        f"{nombre_firma}\t\t\t\t\t\t\t{cpa.nombre}"
    )
    elaborado = doc.add_paragraph(
        f"Elaborado\t\t\t\t\t\t\t\t\tCédula de identidad {cpa.cedula}"
    )
    propietario = doc.add_paragraph(
        f"Propietario\t\t\t\t\t\t\t\t\tContador Público Autorizado N° {cpa.numero_cpa}"
    )
    for p in (firma, elaborado, propietario):
        apply_paragraph_style(
            p, font_name="Arial", font_size=8, bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, line_spacing=1.5
        )
        p.paragraph_format.space_after = Pt(0)

    # 8) Salto de página
    doc.add_page_break()
