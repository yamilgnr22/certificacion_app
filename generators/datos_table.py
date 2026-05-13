# generators/datos_table.py

import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt
from word_helpers import (
    apply_paragraph_style,
    set_row_height,
    set_vertical_alignment,
)

def generar_tabla_datos(doc, df_datos: pd.DataFrame):
    """
    Inserta en `doc` la tabla con la hoja 'Datos' y añade al final
    un título centrado para las fotografías del negocio.
    """
    # 1) Copiar y normalizar DataFrame
    df = df_datos.copy().fillna("")
    num_rows, num_cols = df.shape

    # 2) Crear tabla sin fila de encabezado
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 3) Rellenar celdas con datos y dar formato
    for i, row in df.iterrows():
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cell = cells[j]
            cell.text = str(val)
            apply_paragraph_style(
                cell.paragraphs[0],
                font_name="Arial",
                font_size=8,
                alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
            )
            set_vertical_alignment(cell, "center")
            # Negrita en las dos primeras columnas
            if j < 2:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    # 4) Ancho base para todas las celdas
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(3)

    # 5) Espacio en blanco tras la tabla
    p_gap = doc.add_paragraph()
    p_gap.paragraph_format.space_after = Pt(6)

    # 6) Desactivar autofit y reasignar anchos específicos
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx == 0:
                cell.width = Cm(4)
            elif idx == 1:
                cell.width = Cm(1)
            elif idx == 2:
                cell.width = Cm(13.5)

    # 7) Altura fija de fila (1.0 cm)
    twips = int(1.0 * 1440 / 2.54)
    for row in table.rows:
        set_row_height(row, twips)

    # 8) Salto de página y título de fotografías
    # doc.add_page_break()
    # titulo = doc.add_paragraph("Fotografías del Negocio")
    # apply_paragraph_style(
    #    titulo,
    #    font_name="Arial",
    #    font_size=12,
    #    bold=True,
    #    alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
    #)
