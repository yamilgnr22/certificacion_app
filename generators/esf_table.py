# generators/esf_table.py
"""
Genera la tabla del Estado de Situación Financiera (ESF).

‣ Anchos de columna (cm):   0 → 5.5 | 1 → 2.5 | 2 → 1.0 | 3 → 5.5 | 4 → 2.5
‣ «Unnamed …» en encabezados se deja vacío
‣ Números con separador de miles; negativos entre paréntesis + rojo
‣ Bordes eliminados:
      · fila 0 (encabezado)   · fila 1   · toda la columna 2
"""

import calendar
from datetime import datetime

import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

from config_cpa import load_cpa_profile
from word_helpers import (
    apply_paragraph_style,
    set_row_height,
    set_vertical_alignment,
    set_cell_border,
)
from .utils import extract_cert_fields

# ───────────────────────────── helpers ──────────────────────────────
MESES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _periodo(inicio: datetime, fin: datetime) -> str:
    último = calendar.monthrange(fin.year, fin.month)[1]
    return (
        f"Para el periodo comprendido del 1ro de {MESES[inicio.month-1]} del {inicio.year} "
        f"al {último} de {MESES[fin.month-1]} del {fin.year}"
    )


# ────────────────────────── generador ESF ───────────────────────────
def generar_tabla_esf(doc, df_esf: pd.DataFrame, df_cert: pd.DataFrame) -> None:
    # 1) encabezados (robustos)
    cert      = extract_cert_fields(df_cert)
    nombre    = cert.get("nombre")
    apellido  = cert.get("apellido")
    nombre_completo = (cert.get("nombre_completo") or f"{nombre or ''} {apellido or ''}").strip()
    cedula    = cert.get("cedula")
    direccion = cert.get("direccion_negocio")
    inicio    = cert.get("inicio")
    fin       = cert.get("fin")
    periodo   = _periodo(inicio, fin) if (inicio is not None and fin is not None) else ""

    def _add_center(txt, size=8, bold=False):
        p = doc.add_paragraph(txt)
        apply_paragraph_style(
            p, font_name="Arial", font_size=size, bold=bold,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, line_spacing=1
        )

    _add_center(f"{nombre_completo}", 12, True)
    _add_center(str(cedula), 12, True)
    _add_center(direccion, 8)
    _add_center("Estado de Situación Financiera", 8)
    _add_center("Expresado en córdobas", 8)
    _add_center(periodo, 8)
    doc.add_paragraph()                                # pequeño espacio

    # 2) preparar dataframe
    df = df_esf.copy().fillna("")
    if df.shape[1] < 5:
        raise ValueError(
            "La hoja ESF_Corte debe tener al menos 5 columnas: "
            "descripcion activo, valor activo, separador, descripcion pasivo/patrimonio, valor."
        )
    if df.shape[1] > 5:
        # El formato al corte usa solo las primeras 5 columnas. Algunas
        # plantillas conservan columnas auxiliares/Unnamed a la derecha.
        df = df.iloc[:, :5].copy()
    for col in df.columns[1:]:
        df[col] = df[col].apply(
            lambda x: f"{x:,.0f}" if isinstance(x, (int, float)) else x
        )

    # 3) crear tabla
    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit   = False

    # 3a) encabezado
    for j, col in enumerate(df.columns):
        text = "" if str(col).startswith("Unnamed") else str(col)
        cell = table.rows[0].cells[j]
        cell.text = text
        apply_paragraph_style(
            cell.paragraphs[0], font_name="Arial", font_size=7, bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, line_spacing=1
        )
        set_vertical_alignment(cell, "center")
        set_cell_border(
            cell,
            top={"sz": "3", "val": "single", "color": "000000"},
            bottom={"sz": "3", "val": "single", "color": "000000"},
        )

    # etiquetas de títulos/agrupadores
    encabezados_bold = {
        "Activos", "Pasivos", "Patrimonio",
        "Corrientes", "No Corrientes",
        "Total Corrientes", "Total No Corrientes",
        "Propiedad Planta y Equipos",
        "Total Activos", "Total Pasivos",
        "Total Patrimonio", "Total Pasivo + Patrimonio",
    }
    # totales cuyo VALOR debe ir en negrita
    totales_valor = {
        "Total Corrientes", "Total No Corrientes",
        "Total Activos", "Total Pasivos",
        "Total Patrimonio", "Total Pasivo + Patrimonio",
        "Patrimonio",
    }

    # 3b) datos
    for i, row in df.iterrows():
        cells = table.rows[i + 1].cells
        val0 = str(row.iloc[0]).strip()
        val3 = str(row.iloc[3]).strip() if 3 < len(row) else ""
        es_titulo_izq = val0 in encabezados_bold
        es_titulo_der = val3 in encabezados_bold
        es_total_izq  = val0 in totales_valor
        es_total_der  = val3 in totales_valor

        for j, val in enumerate(row):
            cell = cells[j]
            text = str(val)
            # números negativos con paréntesis y rojo
            if j in (1, 4) and "-" in text:
                text = f"({text.replace('-', '')})"
                cell.text = text
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            else:
                cell.text = text

            # formato
            if j == 0:                                   # descripción izquierda
                apply_paragraph_style(
                    cell.paragraphs[0], font_name="Arial", font_size=7,
                    bold=es_titulo_izq,
                    alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
                    indent=0 if es_titulo_izq else 0.15,
                    line_spacing=1
                )
            elif j == 3:                                 # descripción derecha
                apply_paragraph_style(
                    cell.paragraphs[0], font_name="Arial", font_size=7,
                    bold=es_titulo_der,
                    alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
                    indent=0 if es_titulo_der else 0.15,
                    line_spacing=1
                )
            elif j == 1:                                 # valor izquierda
                apply_paragraph_style(
                    cell.paragraphs[0], font_name="Arial", font_size=7,
                    bold=es_total_izq,
                    alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, line_spacing=1
                )
            elif j == 4:                                 # valor derecha
                apply_paragraph_style(
                    cell.paragraphs[0], font_name="Arial", font_size=7,
                    bold=es_total_der,
                    alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, line_spacing=1
                )
            else:                                        # columna 2 (en blanco)
                apply_paragraph_style(
                    cell.paragraphs[0], font_name="Arial", font_size=7,
                    alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, line_spacing=1
                )

            set_vertical_alignment(cell, "center")
        set_row_height(table.rows[i + 1], 283)           # ≈0.50 cm

    # 3c) bordes para totales (izquierda)
    totales = totales_valor
    for row in table.rows[1:]:
        primer = row.cells[0].text.strip()
        if primer in totales:
            for c in row.cells:
                set_cell_border(c, top={"sz": "3", "val": "single", "color": "000000"})
            if primer in ("Total Activos", "Total Pasivo + Patrimonio"):
                for c in row.cells:
                    set_cell_border(c, bottom={"sz": "3", "val": "double", "color": "000000"})

    # ── ▶︎ NUEVO: bordes especiales para “Total Pasivos” en la columna 3 ◀︎ ──
    for row in table.rows[1:]:
        if row.cells[3].text.strip() == "Total Pasivos":
            for c in (row.cells[3], row.cells[4]):       # descripción y valor
                set_cell_border(
                    c,
                    top={"sz": "3", "val": "single", "color": "000000"},
                    bottom={"sz": "3", "val": "single", "color": "000000"},
                )

    # 4) anchos
    widths = [Cm(5.5), Cm(2.5), Cm(1.0), Cm(5.5), Cm(2.5)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    # ─────────── eliminación de líneas solicitada ────────────
    # filas 0 y 1
    for ridx in (0, 1):
        for c in table.rows[ridx].cells:
            set_cell_border(
                c,
                top={"val": "none"}, bottom={"val": "none"},
                left={"val": "none"}, right={"val": "none"},
            )
    # columna 2 sin bordes
    for r in table.rows:
        set_cell_border(
            r.cells[2],
            top={"val": "none"}, bottom={"val": "none"},
            left={"val": "none"}, right={"val": "none"},
        )

    # 5) firmas y salto
    doc.add_paragraph()
    for _ in range(4):
        p_blank = doc.add_paragraph("")
        p_blank.paragraph_format.line_spacing = 1.5
        p_blank.paragraph_format.space_after  = Pt(0)

    cpa = load_cpa_profile()
    nombre_firma = nombre_completo or (f"{nombre or ''} {apellido or ''}".strip()) or ""
    firmas = (
        f"{nombre_firma}\t\t\t\t\t\t\t{cpa.nombre}",
        f"Elaborado\t\t\t\t\t\t\t\t\tCédula de identidad {cpa.cedula}",
        f"Propietario\t\t\t\t\t\t\t\t\tContador Público Autorizado N° {cpa.numero_cpa}",
    )
    for line in firmas:
        p = doc.add_paragraph(line)
        apply_paragraph_style(
            p, font_name="Arial", font_size=8, bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, line_spacing=1.5
        )
        p.paragraph_format.space_after = Pt(0)

    doc.add_page_break()
