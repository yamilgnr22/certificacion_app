"""
Genera la tabla del Estado de Situación Financiera (ESF) en variante Mensual.

Estructura esperada de la hoja:
  - Columna 0: Descripción (Activos/Pasivos/Patrimonio, subtítulos y cuentas)
  - Columnas 1..N: valores por mes

Formato en el DOCX replica estilos del generador al corte:
  - Encabezado con nombre/cedula/dirección/título/periodo
  - Títulos en negrita
  - Valores con separador de miles; negativos entre paréntesis en rojo (si vienen así del Excel)
"""

from __future__ import annotations

import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import unicodedata

from .utils import extract_cert_fields
from config_cpa import load_cpa_profile
from word_helpers import apply_paragraph_style, set_vertical_alignment, set_cell_border, set_row_height


def _add_header(doc, cert: dict, periodo: str) -> None:
    def add_center(text: str, size=8, bold=False):
        p = doc.add_paragraph(text or "")
        apply_paragraph_style(
            p, font_name="Arial", font_size=size, bold=bold,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, line_spacing=1
        )

    nombre = cert.get("nombre")
    apellido = cert.get("apellido")
    nombre_completo = (cert.get("nombre_completo") or f"{nombre or ''} {apellido or ''}").strip()
    cedula = cert.get("cedula")
    direccion = cert.get("direccion_negocio")

    add_center(nombre_completo, 12, True)
    add_center(str(cedula or ""), 12, True)
    add_center(direccion or "", 8)
    add_center("Estado de Situación Financiera", 8)
    add_center("Expresado en córdobas", 8)
    add_center(periodo or "", 8)
    doc.add_paragraph()  # espacio pequeño


def _periodo_text(cert: dict) -> str:
    inicio = cert.get("inicio")
    fin = cert.get("fin")
    if not inicio or not fin:
        return ""
    import calendar
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    ultimo = calendar.monthrange(fin.year, fin.month)[1]
    return (
        f"Para el periodo comprendido del 1ro de {meses[inicio.month-1]} del {inicio.year} "
        f"al {ultimo} de {meses[fin.month-1]} del {fin.year}"
    )


def _fmt_number(x) -> str:
    try:
        import pandas as pd  # detección de NaN
        if pd.isna(x):
            return ""
        if isinstance(x, (int, float)):
            return f"{x:,.0f}"
        # si viene como string numérico, dejarlo tal cual
        s = str(x)
        return "" if s.lower() == "nan" else s
    except Exception:
        return ""


def generar_tabla_esf_mensual(doc, df_esf_m: pd.DataFrame, df_cert: pd.DataFrame) -> None:
    if df_esf_m is None or df_esf_m.empty:
        # Nada que mostrar; dejar una nota
        p = doc.add_paragraph("ESF Mensual no disponible")
        apply_paragraph_style(p, font_name="Arial", font_size=8, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        return

    cert = extract_cert_fields(df_cert)
    nombre = cert.get("nombre")
    apellido = cert.get("apellido")
    nombre_completo = (cert.get("nombre_completo") or f"{nombre or ''} {apellido or ''}").strip()
    periodo = _periodo_text(cert)
    _add_header(doc, cert, periodo)

    df = df_esf_m.copy()
    # Normalizar encabezados "Unnamed" -> vacío
    cols = ["" if str(c).startswith("Unnamed") else str(c) for c in df.columns]
    df.columns = cols
    # Limpiar NaN en primera columna (descripciones)
    df[df.columns[0]] = df[df.columns[0]].fillna("").astype(str).replace("nan", "")
    # Formatear números como texto con separador de miles
    # Evita FutureWarning de pandas convirtiendo antes a 'object'
    for j in range(1, df.shape[1]):
        colname = df.columns[j]
        df[colname] = df[colname].astype("object").map(_fmt_number)

    # Crear tabla (agrega fila de encabezado)
    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False  # fijar anchos manualmente para parecerse al formato de Excel

    # Limpia todos los bordes por defecto (tabla sin cuadrícula visible)
    for i in range(rows + 1):
        for j in range(cols):
            set_cell_border(table.rows[i].cells[j], top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})

    # Establecer anchos: primera columna amplia, resto estrechas y homogéneas
    # Ajuste solicitado: primera columna 7 cm, resto 1.8 cm
    first_w = Cm(7.0)
    other_w = Cm(1.8)
    widths = [first_w] + [other_w] * (cols - 1)
    for j, w in enumerate(widths):
        for i in range(rows + 1):
            try:
                table.cell(i, j).width = w
            except Exception:
                pass
    # Refuerza anchos por columna para que Word no reasigne el £ltimo
    for j, col in enumerate(table.columns):
        try:
            col.width = widths[j]
        except Exception:
            pass
        for cell in col.cells:
            cell.width = widths[j]

    # Encabezado
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        # Formato de meses (mmm-yy) si es fecha
        header_text = str(col)
        if j >= 1:
            try:
                ts = pd.to_datetime(col, errors="coerce")
                if pd.notna(ts):
                    meses_cortos = ["ene", "feb", "mar", "abr", "may", "jun", "jul", "ago", "sept", "oct", "nov", "dic"]
                    header_text = f"{meses_cortos[ts.month-1]}-{str(ts.year)[2:]}"
            except Exception:
                pass
        cell.text = header_text
        apply_paragraph_style(
            cell.paragraphs[0], font_name="Arial", font_size=7, bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.LEFT if j == 0 else WD_PARAGRAPH_ALIGNMENT.CENTER,
            line_spacing=1
        )
        set_vertical_alignment(cell, "center")
        set_cell_border(cell, top={"sz": "3", "val": "single", "color": "000000"}, bottom={"sz": "3", "val": "single", "color": "000000"})

    encabezados_bold = {
        "Activos", "Pasivos", "Patrimonio",
        "Corrientes", "No Corrientes",
        "Total Corrientes", "Total No Corrientes",
        "Propiedad Planta y Equipos",
        "Total Activos", "Total Pasivos",
        "Total Patrimonio", "Total Pasivo + Patrimonio",
    }
    totales_valor = {
        "Total Corrientes", "Total No Corrientes",
        "Total Activos", "Total Pasivos",
        "Total Patrimonio", "Total Pasivo + Patrimonio",
        "Patrimonio",
    }

    def _shade_row(r_idx: int, fill_hex: str):
        for j in range(cols):
            c = table.rows[r_idx].cells[j]
            tcPr = c._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill_hex)
            tcPr.append(shd)

    current_section = "Activos"

    # Rellenar filas
    for i in range(rows):
        # Detectar fila de depreciaciÃ³n acumulada para formateo negativo
        raw_label = str(df.iloc[i, 0]).strip()
        norm_label = "".join(ch for ch in unicodedata.normalize("NFD", raw_label.lower()) if unicodedata.category(ch) != "Mn")
        is_depreciacion = "depreciacion acumulada" in norm_label or "depreciacion" in norm_label

        for j in range(cols):
            cell = table.rows[i + 1].cells[j]
            val = df.iloc[i, j]
            # Formato especial para la fila de depreciaciÃ³n acumulada: parÃ©ntesis y rojo
            if is_depreciacion and j > 0:
                try:
                    num = float(str(val).replace(",", "").replace("(", "-").replace(")", ""))
                    cell.text = f"({abs(num):,.0f})" if val != "" else ""
                except Exception:
                    cell.text = str(val)
            else:
                cell.text = str(val)
            # estilo
            if j == 0:
                # descripciones, negrita si es encabezado
                bold = str(val).strip() in encabezados_bold
                p = apply_paragraph_style(cell.paragraphs[0], font_name="Arial", font_size=7, bold=bold, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, line_spacing=1)
                # Indentación para cuentas (no encabezados ni totales)
                label = str(val).strip()
                if (not bold) and (not label.lower().startswith("total ")) and label:
                    try:
                        p.paragraph_format.left_indent = Cm(0.6)
                    except Exception:
                        pass
                # Cambiar sección cuando corresponde
                if label.lower() == "pasivos":
                    current_section = "Pasivos"
                elif label.lower() == "patrimonio":
                    current_section = "Patrimonio"
            else:
                p = cell.paragraphs[0]
                apply_paragraph_style(p, font_name="Arial", font_size=7, bold=False, alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, line_spacing=1)
                # valores en negrita si es fila de totales (incluye variante singular)
                if str(df.iloc[i, 0]).strip() in totales_valor or str(df.iloc[i, 0]).strip() == "Total No Corriente":
                    for r in p.runs:
                        r.bold = True
                # negativos entre paréntesis → rojo (incluye fila de depreciación)
                txt = str(cell.text)
                if ((txt.startswith("(") and txt.endswith(")")) or is_depreciacion) and p.runs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(192, 0, 0)

        # Bordes y sombreado por filas clave
        label0 = str(df.iloc[i, 0]).strip()
        top_border = {"sz": "3", "val": "single", "color": "000000"}
        bottom_border = {"sz": "3", "val": "single", "color": "000000"}

        def _apply_row_borders(r_idx: int, *, top: bool = False, bottom: bool = False):
            for j in range(cols):
                set_cell_border(
                    table.rows[r_idx].cells[j],
                    top=top_border if top else None,
                    bottom=bottom_border if bottom else None,
                )

        if label0 == "Total Corrientes":
            _apply_row_borders(i + 1, top=True, bottom=False)
        if label0 in {"Total No Corrientes", "Total No Corriente"}:
            _apply_row_borders(i + 1, top=True, bottom=False)
        if label0 == "Total Activos":
            _apply_row_borders(i + 1, top=True, bottom=True)
        if label0 == "Total Pasivos":
            _apply_row_borders(i + 1, top=True, bottom=True)
        if label0 == "Total Patrimonio":
            _apply_row_borders(i + 1, top=True, bottom=False)
        if label0 == "Total Pasivo + Patrimonio":
            _apply_row_borders(i + 1, top=True, bottom=True)

        # Sombreado para "Total Corrientes" dentro de Pasivos (como imagen)
        if current_section == "Pasivos" and label0 == "Total Corrientes":
            # Se elimina sombreado según requerimiento
            pass
        # Línea superior fina al inicio de cada sección (Pasivos, Patrimonio)
        if label0 in {"Pasivos", "Patrimonio"}:
            for j in range(cols):
                set_cell_border(table.rows[i + 1].cells[j], top=None)

    # Remarcar separadores finales: borde inferior de "Total Pasivo + Patrimonio"
    try:
        # Buscar fila "Total Pasivo + Patrimonio"
        idx = None
        for r in range(rows):
            if str(df.iloc[r, 0]).strip() == "Total Pasivo + Patrimonio":
                idx = r
                break
        if idx is not None:
            for j in range(cols):
                set_cell_border(table.rows[idx + 1].cells[j], bottom={"sz": "3", "val": "single", "color": "000000"})
    except Exception:
        pass

    # Altura uniforme de filas: 0.35 cm
    row_height_twips = int(0.35 * 1440 / 2.54)
    for row in table.rows:
        set_row_height(row, row_height_twips)

    # Espaciado y bloque de firmas (consistente con ESF corte / ER)
    doc.add_paragraph()
    for _ in range(2):
        p_blank = doc.add_paragraph("")
        p_blank.paragraph_format.line_spacing = 1.5
        p_blank.paragraph_format.space_after = Pt(0)

    cpa = load_cpa_profile()
    firma_nombre = nombre_completo or (f"{nombre or ''} {apellido or ''}".strip()) or ""
    firmas = (
        f"{firma_nombre}\t\t\t\t\t\t\t{cpa.nombre_plano}",
        f"Elaborado\t\t\t\t\t\t\t\t\tCedula de identidad {cpa.cedula}",
        f"Propietario\t\t\t\t\t\t\t\t\tContador Publico Autorizado No. {cpa.numero_cpa}",
    )
    for line in firmas:
        p = doc.add_paragraph(line)
        apply_paragraph_style(
            p,
            font_name="Arial",
            font_size=8,
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
            line_spacing=1.5,
        )
        p.paragraph_format.space_after = Pt(0)

    doc.add_page_break()
