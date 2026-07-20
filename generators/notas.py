# generators/notas.py
"""
Notas integradoras (anexos firmados) para el DOCX del motor V2.
Renderiza la lista de motor.notas.Nota con el estilo Arial de la app.
"""

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt

from word_helpers import apply_paragraph_style, set_vertical_alignment


def _fmt(valor) -> str:
    if isinstance(valor, (int, float)):
        return f"{valor:,.0f}"
    return str(valor)


def generar_notas(doc, notas) -> None:
    """Inserta las notas integradoras en una pagina nueva del documento."""
    doc.add_page_break()
    titulo = doc.add_paragraph("Notas a los Estados Financieros (Integraciones)")
    apply_paragraph_style(
        titulo, font_name="Arial", font_size=12, bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, line_spacing=1.15,
    )
    doc.add_paragraph()

    for nota in notas:
        encabezado = doc.add_paragraph(f"{nota.numero}. {nota.titulo}")
        apply_paragraph_style(encabezado, font_name="Arial", font_size=10, bold=True)

        num_cols = len(nota.columnas)
        table = doc.add_table(rows=1 + len(nota.filas) + 1, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"

        # Encabezado
        for j, col in enumerate(nota.columnas):
            cell = table.rows[0].cells[j]
            cell.text = str(col)
            apply_paragraph_style(
                cell.paragraphs[0], font_name="Arial", font_size=8, bold=True,
                alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
            )
            set_vertical_alignment(cell, "center")

        # Filas de detalle
        for i, fila in enumerate(nota.filas, start=1):
            for j, valor in enumerate(fila):
                cell = table.rows[i].cells[j]
                cell.text = _fmt(valor)
                apply_paragraph_style(
                    cell.paragraphs[0], font_name="Arial", font_size=8,
                    alignment=(WD_PARAGRAPH_ALIGNMENT.LEFT if j == 0 else WD_PARAGRAPH_ALIGNMENT.RIGHT),
                )
                set_vertical_alignment(cell, "center")

        # Fila de total (negrita)
        fila_total = table.rows[1 + len(nota.filas)]
        for j, valor in enumerate(nota.total):
            cell = fila_total.cells[j]
            cell.text = _fmt(valor)
            apply_paragraph_style(
                cell.paragraphs[0], font_name="Arial", font_size=8, bold=True,
                alignment=(WD_PARAGRAPH_ALIGNMENT.LEFT if j == 0 else WD_PARAGRAPH_ALIGNMENT.RIGHT),
            )
            set_vertical_alignment(cell, "center")

        # Anchos: descripcion ancha, montos angostos
        table.autofit = False
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = Cm(8) if idx == 0 else Cm(3.8)

        gap = doc.add_paragraph()
        gap.paragraph_format.space_after = Pt(10)

    # La seccion siguiente (Datos del cliente) arranca en pagina propia.
    doc.add_page_break()
