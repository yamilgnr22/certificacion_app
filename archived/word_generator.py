"""word_generator.py
Construye el documento Word a partir de los DataFrames.
Se apoya en las utilidades de `word_helpers`.
"""

from datetime import datetime
import calendar
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from num2words import num2words

from word_helpers import (
    apply_paragraph_style,
    add_paragraph_border,
    set_row_height,
    set_vertical_alignment,
    set_cell_border,
    add_page_number,
)


class DocxGenerator:
    _SPANISH_MONTHS_ABBR = {
        1: "Ene", 2: "Feb", 3: "Mar", 4: "Abr",
        5: "May", 6: "Jun", 7: "Jul", 8: "Ago",
        9: "Sep", 10: "Oct", 11: "Nov", 12: "Dic",
    }

    def __init__(self, df_esf, df_er, df_datos, df_certificacion):
        self.df_esf = df_esf
        self.df_er = df_er
        self.df_datos = df_datos
        self.df_certificacion = df_certificacion
        self.doc = Document()
        self._extract_certificacion_data()

    # ------------------------------------------------------------------ #
    # DATOS CERTIFICACIÓN                                                 #
    # ------------------------------------------------------------------ #
    def _extract_certificacion_data(self):
        dfc = self.df_certificacion
        self.nombre_cliente = dfc.iloc[0, 1]
        self.apellido_cliente = dfc.iloc[1, 1]
        self.cedula = dfc.iloc[2, 1]
        self.periodo_inicio = pd.to_datetime(dfc.iloc[3, 1])
        self.periodo_fin = pd.to_datetime(dfc.iloc[4, 1])
        self.estado_civil = dfc.iloc[5, 1]
        self.profesion = dfc.iloc[6, 1]
        self.sexo = dfc.iloc[7, 1]
        self.domicilio = dfc.iloc[8, 1]
        self.direccion_personal = dfc.iloc[9, 1]
        self.direccion_negocio = dfc.iloc[10, 1]
        self.primer_apellido = dfc.iloc[11, 1]
        self.ingresos_brutos = dfc.iloc[12, 1]
        self.ingresos_promedio = dfc.iloc[13, 1]
        self.utilidad_periodo = dfc.iloc[14, 1]
        self.utilidad_promedio = dfc.iloc[15, 1]
        self.banco = dfc.iloc[16, 1]
        self.fecha_certificacion = pd.to_datetime(dfc.iloc[17, 1])

        # Convertir números a palabras
        self.ingresos_brutos_palabras = num2words(self.ingresos_brutos, lang="es")
        self.ingresos_promedio_palabras = num2words(self.ingresos_promedio, lang="es")
        self.utilidad_periodo_palabras = num2words(self.utilidad_periodo, lang="es")
        self.utilidad_promedio_palabras = num2words(self.utilidad_promedio, lang="es")

        self.fecha_certificacion_palabras = self._fecha_a_palabras(self.fecha_certificacion)

        # Números formateados
        self.ingresos_brutos_fmt = f"{self.ingresos_brutos:,.0f}"
        self.ingresos_promedio_fmt = f"{self.ingresos_promedio:,.0f}"
        self.utilidad_periodo_fmt = f"{self.utilidad_periodo:,.0f}"
        self.utilidad_promedio_fmt = f"{self.utilidad_promedio:,.0f}"


    # ------------------------------------------------------------------ #
    # PROPIEDADES DERIVADAS DEL SEXO                                     #
    # ------------------------------------------------------------------ #
    @property
    def genero_cliente(self) -> str:
        """
        Devuelve 'la señora' si self.sexo es 'Femenino', 
        en cualquier otro caso 'el señor'.
        """
        if isinstance(self.sexo, str) and self.sexo.strip().lower() == "femenino":
            return "la señora"
        return "el señor"

    @property
    def genero_cliente_2(self) -> str:
        """
        Igual que arriba pero formato 'de la señora' / 'del señor'.
        """
        if isinstance(self.sexo, str) and self.sexo.strip().lower() == "femenino":
            return "de la señora"
        return "del señor"





    # ------------------------------------------------------------------ #
    # UTILIDADES DE FECHA                                                #
    # ------------------------------------------------------------------ #
    @staticmethod
    def _fecha_a_palabras(fecha):
        meses = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
        ]
        dia = num2words(fecha.day, lang="es")
        mes = meses[fecha.month - 1]
        anio = num2words(fecha.year, lang="es")
        return f"{dia} días del mes de {mes} del año {anio}"

    def _periodo_certificacion(self):
        meses_es = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
        ]
        last_day_fin = calendar.monthrange(self.periodo_fin.year, self.periodo_fin.month)[1]
        return (
            f"Para el periodo comprendido del 1ro de {meses_es[self.periodo_inicio.month - 1]} del año {self.periodo_inicio.year} "
            f"al {last_day_fin} de {meses_es[self.periodo_fin.month - 1]} del año {self.periodo_fin.year}"
        )

    # ------------------------------------------------------------------ #
    # ATALHOS DE PÁRRAFOS                                                #
    # ------------------------------------------------------------------ #
    def _add_center(self, text, size=8, bold=False):
        p = self.doc.add_paragraph(text)
        apply_paragraph_style(
            p,
            font_name="Arial",
            font_size=size,
            bold=bold,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
            line_spacing=1,
        )

    def _add_justify(self, text, size=7, bold=False):
        p = self.doc.add_paragraph(text)
        apply_paragraph_style(
            p,
            font_name="Arial",
            font_size=size,
            bold=bold,
            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            line_spacing=1.5,
        )

    # ------------------------------------------------------------------ #
    # ENCABEZADO Y PIE                                                   #
    # ------------------------------------------------------------------ #
    def _configurar_encabezado_pie(self):
        for section in self.doc.sections:
            # Márgenes
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)

            # ---------- Encabezado ----------
            header = section.header
            if header.paragraphs:
                p_header = header.paragraphs[0]
                p_header.clear()
            else:
                p_header = header.add_paragraph()

            run1 = p_header.add_run("Licenciado Yamil René García Laguna\n")
            run1.bold = True
            run1.font.size = Pt(10)
            run1.font.name = "Abadi"
            run1.font.color.rgb = RGBColor(20, 47, 80)

            run2 = p_header.add_run("Contador Público Autorizado No. 3314")
            run2.font.name = "Abadi Extra Light"
            run2.font.size = Pt(9)

            apply_paragraph_style(p_header, font_name="Arial", font_size=9)
            add_paragraph_border(p_header, "bottom")
            
              # <-- Aquí añado un espacio posterior de 12 pt al encabezado
            p_header.paragraph_format.space_after = Pt(12)

            # ---------- Pie de página ----------
            footer = section.footer
            if footer.paragraphs:
                p_footer = footer.paragraphs[0]
                p_footer.clear()
            else:
                p_footer = footer.add_paragraph()

            p_footer.add_run(
                "📞 +505 8966 5057   📧 yamilgnr22@gmail.com\t\tPágina "
            )
            
            # Aplicar estilo al texto existente
            apply_paragraph_style(p_footer, font_name="Abadi Extra Light", font_size=8)
            add_paragraph_border(p_footer, "top")
            
            # Insertar número de página
            add_page_number(p_footer)
            
            for run in p_footer.runs:
                run.font.name = "Abadi Extra Light"
                run.font.size = Pt(8)


    # ------------------------------------------------------------------ #
    # CERTIFICACIÓN                                                      #
    # ------------------------------------------------------------------ #
    from docx.shared import Pt  # normalmente ya lo tienes importado

    def _generar_certificacion(self):
        self._add_center("CERTIFICACIÓN DEL CONTADOR PÚBLICO INDEPENDIENTE", 8, True)
        # —> Espacio posterior de 12 pt al título
        titulo_parrafo = self.doc.paragraphs[-1]
        titulo_parrafo.paragraph_format.space_after = Pt(12)
        periodo = self._periodo_certificacion()

        cuerpo = [
            f"Yo, Yamil René García Laguna, mayor de edad, soltero, Licenciado en Contaduría Pública y Auditoría, del domicilio de Managua, identificado con cédula de identidad número 001-281186-0054R, en mi calidad de Contador Público Autorizado para ejercer la profesión por el excelentísimo Ministerio de Educación Cultura y Deporte, bajo acuerdo C.P.A. No. 315-2023, fechado 22 de diciembre del 2023, por el quinquenio que finalizará el 21 de diciembre del 2028; expreso a través de este documento, que conforme a las leyes vigentes del país, en mi carácter de profesional de Contaduría Pública y en representación propia, que fui contratado para Certificar el Estado de Situación Financiera y el Estado de Resultados {periodo}, preparados por {self.genero_cliente} {self.nombre_cliente} {self.apellido_cliente}, mayor de edad, {self.estado_civil}, {self.profesion}, quien se identifica con la cédula de identidad No. {self.cedula}, con domicilio en {self.domicilio}.",
            f"Responsabilidad: {self.genero_cliente} {self.nombre_cliente} {self.apellido_cliente}, es responsable sobre la información suministrada y reflejada en los estados financieros presentados para la Certificación, mi responsabilidad consiste en Certificar que las cifras contenidas en dichos estados financieros están conformes con los registros contables llevados {self.genero_cliente} {self.nombre_cliente} {self.apellido_cliente}.",
            f"Mi trabajo fue realizado de acuerdo con la Normativa sobre “Trabajo previamente convenidos”, y los objetivos que se persiguieron con dicho trabajo fueron los siguientes:",
            # Estos tres párrafos se mostrarán con viñeta
            f"Determinar que los estados financieros presentados por {self.genero_cliente} {self.primer_apellido} fueron preparados de acuerdo con los registros contables llevados para registrar las operaciones.",
            f"Determinar que los registros contables contenidos en los estados financieros {self.genero_cliente_2} {self.primer_apellido} se encuentran de acuerdo con principios de contabilidad generalmente aceptados en Nicaragua.",
            f"Determinar si los ingresos netos derivados de la actividad económica {self.genero_cliente_2} {self.primer_apellido} se encuentran presentados de forma razonable de conformidad con sus registros contables.",
            f"Para lograr los objetivos efectué una revisión selectiva de los registros contables a fin de determinar que estos estaban efectuados de acuerdo con Principios de Contabilidad Generalmente Aceptados en Nicaragua.",
            f"Mi trabajo proporciona una base razonable para Certificar que las cifras contenidas en los Estados Financieros {self.genero_cliente_2} {self.nombre_cliente} {self.apellido_cliente}, han sido preparados de acuerdo con los registros contables de sus operaciones a la fecha anteriormente indicada. En este sentido, sobre la base del trabajo que efectué, Certifico que:",
            # Los siguientes dos párrafos deben tener viñeta también:
            f"• Los ingresos brutos para el periodo revisado ascendieron a NIO {self.ingresos_brutos_fmt} ({self.ingresos_brutos_palabras} córdobas), que da como resultado un promedio mensual de ingresos brutos de NIO {self.ingresos_promedio_fmt} ({self.ingresos_promedio_palabras} córdobas).",
            f"• Las utilidades netas del periodo revisado (después de deducir costos y gastos) fueron de NIO {self.utilidad_periodo_fmt} ({self.utilidad_periodo_palabras} córdobas), que da como resultado un promedio mensual de utilidades netas de NIO {self.utilidad_promedio_fmt} ({self.utilidad_promedio_palabras} córdobas).",
            f"Se adjuntan a esta Certificación, el Estado de Resultados, el Estado de Situación Financiera y los anexos a los estados financieros, los cuales han sido rubricado y sellado por el suscrito Contador Público Autorizado.",
            f"Esta certificación ha sido solicitada para completar los requisitos bancarios con {self.banco}, por lo que no debe ser utilizada para otro trámite legal ante cualquier otra institución pública o privada.",
            f"Dado en la ciudad de Managua, a los {self.fecha_certificacion_palabras}."
        ]
        bullets = {3, 4, 5, 8, 9}

        for idx, texto in enumerate(cuerpo):
            if idx in bullets:
                # 1) Crear párrafo de viñeta
                texto_sin = texto.lstrip("• ")
                p = self.doc.add_paragraph(texto_sin)
                if "List Bullet" in self.doc.styles:
                    p.style = "List Bullet"

                # 2) Aplicar formato base
                apply_paragraph_style(
                    p,
                    font_name="Arial",
                    font_size=7,
                    alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                    line_spacing=1.5
                )

                # 3) Añadir sangría:
                #    - todo el párrafo a 1 cm de la izquierda
                #    - la viñeta (primera línea) retrocede 0.5 cm
                fmt = p.paragraph_format
                fmt.left_indent        = Cm(0.50)
                fmt.first_line_indent  = Cm(-0.50)

            else:
                # Párrafo normal justificado
                p = self.doc.add_paragraph(texto)
                apply_paragraph_style(
                    p,
                    font_name="Arial",
                    font_size=7,
                    alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                    line_spacing=1.5
                )

            # 4) Espacio posterior uniforme (por ejemplo 6 pt)
            p.paragraph_format.space_after = Pt(12)
            
            # ─── Aquí insertamos DOS PÁRRAFOS VACÍOS con estilo idéntico ───
        for _ in range(2):
            p_blank = self.doc.add_paragraph("")
            apply_paragraph_style(
                p_blank,
                font_name="Arial",
                font_size=7,
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                line_spacing=1.5
            )
            p_blank.paragraph_format.space_after = Pt(12)

    
        # Firmas
        for line in [
            "Yamil René García Laguna",
            "Cédula de identidad 001-281186-0054R",
            "Contador Público Autorizado Nº 3314",
        ]:
            p = self.doc.add_paragraph(line)
            apply_paragraph_style(
                p,
                font_name="Arial",
                font_size=7,
                bold=True,
                alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            )
            p.paragraph_format.space_after = Pt(0)

        self.doc.add_page_break()



    # ------------------------------------------------------------------ #
    # TABLA DEL ESTADO DE RESULTADOS (ER)                                #
    # ------------------------------------------------------------------ #
    def generar_tabla_er(self):
        """
        Genera la tabla del Estado de Resultados (ER) y agrega las firmas correspondientes.
        """
        # 1) Encabezados centrados antes de la tabla
        self._add_center(f"{self.nombre_cliente} {self.apellido_cliente}", size=12, bold=True)
        self._add_center(str(self.cedula), size=12, bold=True)
        self._add_center(str(self.direccion_negocio), size=8)
        self._add_center("Estado de Resultados", size=8)
        self._add_center("Expresado en Córdobas", size=8)
        linea_6_text = self._periodo_certificacion()
        self._add_center(linea_6_text, size=8)
        # Un poco de espacio antes de la tabla
        for _ in range(1):
            p = self.doc.add_paragraph(" ")
            apply_paragraph_style(
                p,
                font_name="Arial",
                font_size=7,
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                line_spacing=1.5
            )
            p.paragraph_format.space_after = Pt(6)

        # 2) Preprocesar DataFrame df_er
        df = (
            self.df_er.copy()
            .drop(index=[0, 1, 2, 3, 4])
            .reset_index(drop=True)
            .drop(columns=[self.df_er.columns[1]])
            .fillna("")
        )
        for col in df.columns[1:]:
            df[col] = df[col].apply(lambda x: f"{x:,.0f}" if isinstance(x, (int, float)) else x)

        # 3) Crear tabla con fila extra para encabezados
        num_rows, num_cols = df.shape
        table = self.doc.add_table(rows=num_rows + 1, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # 4) Encabezados de columna
        for j, col_name in enumerate(df.columns):
            cell = table.rows[0].cells[j]
            cell.text = self._format_header_to_spanish(col_name)
            apply_paragraph_style(
                cell.paragraphs[0],
                font_name="Arial",
                font_size=7,
                bold=True,
                alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                line_spacing=1,
                space_before=0,
                space_after=0
            )
            set_vertical_alignment(cell, "center")
            set_cell_border(
                cell,
                top={"sz": "3", "val": "single", "color": "000000"},
                bottom={"sz": "3", "val": "single", "color": "000000"}
            )

        # 5) Rellenar filas de datos
        no_indent = ["Descripción", "Ingresos", "(=) Ingresos Brutos",
                     "(-) Gastos operativos", "Total gastos operativos", "Ingresos/Utilidad Neta"]
        for i, row in df.iterrows():
            cells = table.rows[i + 1].cells
            for j, val in enumerate(row):
                cell = cells[j]
                cell.text = str(val)
                para = cell.paragraphs[0]
                if row.iloc[0] in no_indent:
                    align = WD_PARAGRAPH_ALIGNMENT.RIGHT if j > 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
                    apply_paragraph_style(
                        para, font_name="Arial", font_size=7, bold=True,
                        alignment=align, line_spacing=1, space_before=0, space_after=0
                    )
                else:
                    if j == 0:
                        apply_paragraph_style(
                            para, font_name="Arial", font_size=7,
                            alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
                            indent=0.15, line_spacing=1, space_before=0, space_after=0
                        )
                    else:
                        apply_paragraph_style(
                            para, font_name="Arial", font_size=7,
                            alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT,
                            line_spacing=1, space_before=0, space_after=0
                        )
                set_vertical_alignment(cell, "center")
            set_row_height(table.rows[i + 1], 283)  # ~0.50 cm

        # 6) Bordes especiales
        for i, row in enumerate(table.rows[1:], start=1):
            first = row.cells[0].text.strip()
            if first in ["(=) Ingresos Brutos", "Total gastos operativos"]:
                for cell in row.cells:
                    set_cell_border(cell, top={"sz": "3", "val": "single", "color": "000000"})
        # borde superior simple y borde inferior doble en la última fila
        for cell in table.rows[num_rows].cells:
            set_cell_border(
                cell,
                top={"sz": "3", "val": "single", "color": "000000"},
                bottom={"sz": "3", "val": "double", "color": "000000"}
            )

        # 7) Ajuste de anchos y alturas
        for row in table.rows:
            for k, cell in enumerate(row.cells):
                cell.width = Cm(4) if k == 0 else Cm(1.8)
        # fila 0 → 397 twips (~0.70 cm), resto → 283 twips (~0.50 cm)
        set_row_height(table.rows[0], 397)

        # 8) Espacio y firma al final
        for _ in range(5):
            p = self.doc.add_paragraph(" ")
            apply_paragraph_style(
                p, font_name="Arial", font_size=8,
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, line_spacing=1.5
            )
            p.paragraph_format.space_after = Pt(6)

        firma = self.doc.add_paragraph(f"{self.nombre_cliente} {self.apellido_cliente}        Yamil René García Laguna")
        apply_paragraph_style(firma, font_name="Arial", font_size=8, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        elab = self.doc.add_paragraph("Elaborado        Cédula de identidad 001-281186-0054R")
        apply_paragraph_style(elab, font_name="Arial", font_size=8, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        prop = self.doc.add_paragraph("Propietario      Contador Público Autorizado N° 3314")
        apply_paragraph_style(prop, font_name="Arial", font_size=8, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

        # 9) Salto de página tras la tabla ER
        self.doc.add_page_break()


    # ------------------------------------------------------------------ #
    # DOCUMENTO COMPLETO                                                 #
    # ------------------------------------------------------------------ #
    def generar_documento_completo(self, output_path):
        """Construye el documento final y lo guarda."""
        self._configurar_encabezado_pie()
        self._generar_certificacion()
        self.generar_tabla_er()
        # Más tablas: generar_tabla_esf(), generar_tabla_datos()
        self.doc.save(output_path)

    
    
    
